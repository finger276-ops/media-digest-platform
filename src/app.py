from __future__ import annotations

import argparse
import os
import tempfile
import uuid
import re
import textwrap
from io import BytesIO
from datetime import date, datetime
from pathlib import Path
from typing import Any

import pandas as pd
import streamlit as st
import altair as alt

from import_adapters import read_source_table, get_excel_sheet_names
from io_utils import read_table
from services.cached_store import (
    supabase_configured,
    list_projects,
    create_project,
    update_project,
    resolve_project_access,
    list_periods,
    make_period_id,
    save_processed_tables,
    load_generated_tables,
    update_period_metadata,
    delete_period,
    save_uploaded_file_to_storage,
    get_manual,
    list_manual,
    save_manual,
    delete_manual,
    clear_platform_caches,
)
from preprocess import run_preprocess_from_dataframe
from services.metrics_compute import (
    numeric_series,
    prepare_dashboard_messages,
    format_int,
    sentiment_counts,
    percent_text,
    overview_metrics,
)
from services.tag_compute import (
    split_pipe_values,
    normalize_tag_key,
    declared_ba_tag_set,
    is_brand_analytics_messages,
    clean_brand_analytics_tags,
    build_tag_statistics_compute,
)
from services.message_compute import message_text_column, message_link_column
from services.perf import perf_block, render_perf_sidebar, reset_perf_events

APP_TITLE = "Платформа дайджестов"
APP_VERSION = "4.4.6: unified selected tag view"

ALGORITHM_PROFILE_OPTIONS = {
    "universal": "Универсальный",
    "brand_monitoring": "Бренд-мониторинг",
    "construction_materials": "Строительство / материалы",
    "driver_chats": "Дайджест водительских чатов",
    "taxi_legacy": "Такси / водительские чаты (legacy)",
}

TAXI_PROJECT_PROFILES = {"driver_chats", "taxi_legacy"}

SENTIMENT_COLOR_DOMAIN = ["Позитив", "Нейтрал", "Негатив"]
SENTIMENT_COLOR_RANGE = ["#2ca02c", "#9e9e9e", "#d62728"]

CHART_LABEL_POSITION_OPTIONS = {
    "top": "У верхнего края",
    "center": "В центре",
    "bottom": "У нижнего края",
}
CHART_LABEL_FONT_OPTIONS = ["Arial", "Inter", "Roboto", "Verdana", "Tahoma", "Times New Roman"]
DEFAULT_CHART_LABEL_SETTINGS = {
    "font": "Arial",
    "font_size": 11,
    "position": "top",
    "show_donut_legend": False,
}

REPORT_TEMPLATE_OPTIONS = {
    "summary": "Краткое саммари",
    "client_overview": "Клиентский обзор",
    "comparison": "Сравнительный отчет",
    "full": "Полный отчет",
}

DEFAULT_REPORT_BRANDING = {
    "client_name": "",
    "report_title": "Дайджест упоминаний",
    "accent_color": "#2563eb",
    "background_color": "#ffffff",
    "footer_text": "",
    "logo_url": "",
}


def _valid_hex_color(value: Any, fallback: str) -> str:
    text = str(value or "").strip()
    if re.match(r"^#[0-9a-fA-F]{6}$", text):
        return text
    return fallback


def report_branding_from_project_settings(settings: dict[str, Any] | None, *, project_name: str = "") -> dict[str, Any]:
    raw = {}
    if isinstance(settings, dict):
        raw = settings.get("report_branding") or {}
    if not isinstance(raw, dict):
        raw = {}
    result = dict(DEFAULT_REPORT_BRANDING)
    result.update({k: str(raw.get(k) or result.get(k) or "").strip() for k in ["client_name", "report_title", "footer_text", "logo_url"]})
    result["accent_color"] = _valid_hex_color(raw.get("accent_color"), result["accent_color"])
    result["background_color"] = _valid_hex_color(raw.get("background_color"), result["background_color"])
    if not result.get("client_name"):
        result["client_name"] = str(project_name or "").strip()
    if not result.get("report_title"):
        result["report_title"] = "Дайджест упоминаний"
    return result



def project_topic_profile(project_row: pd.Series | None) -> str:
    settings = project_settings_from_row(project_row) if project_row is not None else {}
    profile = str(settings.get("topic_profile") or "universal")
    return profile if profile in ALGORITHM_PROFILE_OPTIONS else "universal"


def is_taxi_project_profile(profile: str) -> bool:
    return str(profile or "").strip() in TAXI_PROJECT_PROFILES


def is_taxi_project_row(project_row: pd.Series | None) -> bool:
    return is_taxi_project_profile(project_topic_profile(project_row))


def is_brand_analytics_event_set(events: pd.DataFrame) -> bool:
    """Return True when events were built from Brand Analytics `Сюжет` values.

    Brand Analytics projects can legitimately have small one-off сюжеты, so the
    default small-event filter should not hide them. Algorithmic projects
    (especially driver chats) are much noisier and need a higher threshold.
    """
    if events is None or events.empty or "event_source" not in events.columns:
        return False
    values = events["event_source"].fillna("").astype(str).str.lower()
    return bool(values.str.contains("brand_analytics_story", regex=False).any())


def default_min_event_messages(profile: str, events: pd.DataFrame | None = None) -> int:
    """Default threshold for showing information events in dashboards.

    - Brand Analytics: keep every `Сюжет`, because the source system already
      provides editorial/story grouping.
    - Driver chats and other algorithmic projects: hide tiny clusters by default
      so one-message noise does not become an information event.
    """
    if is_brand_analytics_event_set(events if events is not None else pd.DataFrame()):
        return 1
    return 4


def render_min_event_messages_control(profile: str, events: pd.DataFrame | None = None, *, key: str = "min_event_messages") -> int:
    default_value = int(default_min_event_messages(profile, events))
    max_value = 50
    help_text = (
        "Инфоповоды с меньшим числом сообщений скрываются из таблицы и саммари. "
        "Сообщения при этом остаются в общей статистике и полной ленте."
    )
    return int(st.sidebar.number_input(
        "Мин. сообщений в инфоповоде",
        min_value=1,
        max_value=max_value,
        value=default_value,
        step=1,
        help=help_text,
        key=key,
    ))


def filter_small_events(events_agg: pd.DataFrame, min_messages: int) -> tuple[pd.DataFrame, int, int]:
    """Hide tiny information events from dashboard-level analytics.

    This does not delete events or messages from storage; it only filters the
    analytical view. It is intended to suppress one-off algorithmic clusters
    such as `Обсуждение: ...` with 1-3 messages.
    """
    if events_agg is None or events_agg.empty or min_messages <= 1 or "message_count" not in events_agg.columns:
        return events_agg, 0, 0
    work = events_agg.copy()
    counts = pd.to_numeric(work["message_count"], errors="coerce").fillna(0).astype(int)
    keep_mask = counts >= int(min_messages)
    hidden_events = int((~keep_mask).sum())
    hidden_messages = int(counts[~keep_mask].sum())
    return work[keep_mask].copy(), hidden_events, hidden_messages


def render_small_events_notice(hidden_events: int, hidden_messages: int, min_messages: int) -> None:
    if hidden_events <= 0:
        return
    st.caption(
        f"Скрыто малых инфоповодов: {format_int(hidden_events)} "
        f"(< {format_int(min_messages)} сообщений). "
        f"Сообщений в них: {format_int(hidden_messages)}. "
        "Они не удалены и остаются в общей статистике/ленте."
    )


def project_settings_from_row(row) -> dict[str, Any]:
    settings = {}
    try:
        settings = row.get("settings") or {}
    except Exception:
        settings = {}
    return settings if isinstance(settings, dict) else {}


def chart_label_settings_from_project_settings(settings: dict[str, Any] | None) -> dict[str, Any]:
    raw = {}
    if isinstance(settings, dict):
        raw = settings.get("chart_label_settings") or {}
    if not isinstance(raw, dict):
        raw = {}
    result = dict(DEFAULT_CHART_LABEL_SETTINGS)
    font = str(raw.get("font") or result["font"]).strip()
    result["font"] = font or result["font"]
    try:
        size = int(raw.get("font_size", result["font_size"]))
    except Exception:
        size = int(result["font_size"])
    result["font_size"] = max(8, min(28, size))
    position = str(raw.get("position") or result["position"]).strip()
    result["position"] = position if position in CHART_LABEL_POSITION_OPTIONS else result["position"]
    result["show_donut_legend"] = bool(raw.get("show_donut_legend", result.get("show_donut_legend", False)))
    return result


def chart_label_text_kwargs(settings: dict[str, Any] | None, *, chart_type: str = "line") -> dict[str, Any]:
    cfg = chart_label_settings_from_project_settings({"chart_label_settings": settings or {}})
    position = cfg.get("position", "top")
    kwargs: dict[str, Any] = {
        "align": "center",
        "font": cfg.get("font", "Arial"),
        "size": int(cfg.get("font_size", 11)),
    }
    if position == "center":
        kwargs.update({"baseline": "middle", "dy": 0})
    elif position == "bottom":
        kwargs.update({"baseline": "top", "dy": 10})
    else:
        kwargs.update({"baseline": "bottom", "dy": -8})
    if chart_type == "bar" and position == "bottom":
        kwargs.update({"baseline": "bottom", "dy": -4})
    return kwargs


def chart_label_radius(settings: dict[str, Any] | None) -> int:
    cfg = chart_label_settings_from_project_settings({"chart_label_settings": settings or {}})
    position = cfg.get("position", "top")
    if position == "center":
        return 72
    if position == "bottom":
        return 54
    return 104



def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(add_help=False)
    parser.add_argument("--work-dir", default=os.getenv("PLATFORM_WORK_DIR", "data/work"))
    args, _ = parser.parse_known_args()
    return args


def get_secret_value(name: str, default: str = "") -> str:
    try:
        value = st.secrets.get(name, default)
    except Exception:
        value = os.getenv(name, default)
    return str(value or "").strip()


def is_platform_admin() -> bool:
    admin_password = get_secret_value("PLATFORM_ADMIN_PASSWORD") or get_secret_value("ADMIN_PASSWORD")
    if "platform_is_admin" not in st.session_state:
        st.session_state["platform_is_admin"] = False
    if not admin_password:
        st.sidebar.warning("PLATFORM_ADMIN_PASSWORD не настроен: режим владельца временно доступен всем.")
        return True
    if st.session_state.get("platform_is_admin"):
        st.sidebar.success("Режим: владелец платформы")
        if st.sidebar.button("Выйти из режима владельца"):
            st.session_state["platform_is_admin"] = False
            st.rerun()
        return True
    with st.sidebar.expander("Вход владельца платформы", expanded=False):
        password = st.text_input("Пароль владельца", type="password", key="platform_admin_password")
        if st.button("Войти", key="platform_admin_login"):
            if password == admin_password:
                st.session_state["platform_is_admin"] = True
                st.rerun()
            else:
                st.error("Неверный пароль.")
    return False


def role_rank(role: str) -> int:
    return {"none": 0, "viewer": 1, "editor": 2, "owner": 3}.get(role, 0)


def fmt_date(value: Any) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    try:
        ts = pd.to_datetime(value, errors="coerce", dayfirst=True)
        if pd.isna(ts):
            return ""
        return ts.strftime("%d.%m.%Y")
    except Exception:
        return ""


def fmt_period(row: pd.Series) -> str:
    start = fmt_date(row.get("date_from") or row.get("start_date"))
    end = fmt_date(row.get("date_to") or row.get("end_date"))
    if start and end and start != end:
        return f"{start}–{end}"
    return start or end


def normalize_text(value: Any) -> str:
    return str(value or "").strip()


def first_existing_col(df: pd.DataFrame, columns: list[str | None]) -> str | None:
    if df is None or not isinstance(df, pd.DataFrame):
        return None
    for col in columns or []:
        if col and col in df.columns:
            return str(col)
    return None



def render_overview_statistics(messages: pd.DataFrame) -> None:
    """Render top-level numbers for the start page."""
    st.subheader("Статистика")
    total_messages = int(len(messages)) if isinstance(messages, pd.DataFrame) else 0
    audience = int(numeric_series(messages, ["audience", "Аудитория"]).sum()) if total_messages else 0
    reach = int(numeric_series(messages, ["views", "Просмотры", "Просмотров", "reach", "Охват"]).sum()) if total_messages else 0
    engagement = int(numeric_series(messages, ["engagement", "Вовлечённость", "Вовлеченность", "engagement_count"]).sum()) if total_messages else 0
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Сообщений", format_int(total_messages))
    c2.metric("Суммарная аудитория", format_int(audience))
    c3.metric("Суммарный охват", format_int(reach))
    c4.metric("Суммарная вовлеченность", format_int(engagement))


@st.cache_data(show_spinner=False)
def build_tag_statistics(messages: pd.DataFrame) -> pd.DataFrame:
    return build_tag_statistics_compute(messages)

def selected_period_label(periods: pd.DataFrame, period_ids: list[str]) -> str:
    """Human-readable label for the currently selected period set."""
    ids = [str(x) for x in (period_ids or []) if str(x).strip()]
    if not ids:
        return "выбранный период"
    if periods is None or periods.empty or "period_id" not in periods.columns:
        return ", ".join(ids[:3]) + (f" и еще {len(ids) - 3}" if len(ids) > 3 else "")

    subset = periods[periods["period_id"].astype(str).isin(ids)].copy()
    if subset.empty:
        return ", ".join(ids[:3]) + (f" и еще {len(ids) - 3}" if len(ids) > 3 else "")

    if len(subset) == 1:
        row = subset.iloc[0]
        name = str(row.get("period_name") or "").strip()
        period = fmt_period(row)
        return f"{name} · {period}" if name and period else name or period or ids[0]

    dates: list[pd.Timestamp] = []
    for col in ["date_from", "date_to", "start_date", "end_date"]:
        if col in subset.columns:
            parsed = pd.to_datetime(subset[col], errors="coerce", dayfirst=True).dropna()
            dates.extend(parsed.tolist())
    if dates:
        start_s = fmt_date(min(dates))
        end_s = fmt_date(max(dates))
        date_part = f"{start_s}–{end_s}" if start_s and end_s and start_s != end_s else start_s or end_s
    else:
        date_part = ""

    names = [str(x).strip() for x in subset.get("period_name", pd.Series(dtype=str)).fillna("").tolist() if str(x).strip()]
    if len(names) <= 3 and names:
        name_part = "; ".join(names)
    else:
        name_part = f"{len(subset)} период(а/ов)"
    return f"{name_part} · {date_part}" if date_part else name_part


def _period_row_label(row: pd.Series, fallback: str = "") -> str:
    name = str(row.get("period_name") or "").strip() if isinstance(row, pd.Series) else ""
    period = fmt_period(row) if isinstance(row, pd.Series) else ""
    if name and period:
        return f"{name} · {period}"
    return name or period or fallback or "период"


def _selected_period_rows(periods: pd.DataFrame, period_ids: list[str]) -> pd.DataFrame:
    ids = [str(x) for x in (period_ids or []) if str(x).strip()]
    if not ids or periods is None or periods.empty or "period_id" not in periods.columns:
        return pd.DataFrame({"period_id": ids})
    work = periods[periods["period_id"].astype(str).isin(ids)].copy()
    if work.empty:
        return pd.DataFrame({"period_id": ids})
    # Preserve missing selected ids so comparison does not silently lose a period.
    existing = set(work["period_id"].astype(str))
    missing = [pid for pid in ids if pid not in existing]
    if missing:
        work = pd.concat([work, pd.DataFrame({"period_id": missing})], ignore_index=True)
    return work


def _ordered_period_ids(periods: pd.DataFrame, period_ids: list[str]) -> list[str]:
    """Return selected period ids in chronological order for sequence comparison."""
    ids = [str(x) for x in (period_ids or []) if str(x).strip()]
    if len(ids) <= 1:
        return ids
    rows = _selected_period_rows(periods, ids).copy()
    if rows.empty or "period_id" not in rows.columns:
        return ids

    date_series = pd.Series(pd.NaT, index=rows.index, dtype="datetime64[ns]")
    for col in ["date_from", "start_date", "uploaded_at", "date_to", "end_date"]:
        if col in rows.columns:
            parsed = pd.to_datetime(rows[col], errors="coerce", dayfirst=True)
            date_series = date_series.combine_first(parsed)
    rows["_sort_date"] = date_series
    rows["_input_order"] = rows["period_id"].astype(str).map({pid: i for i, pid in enumerate(ids)})
    rows = rows.sort_values(["_sort_date", "_input_order"], na_position="last", kind="mergesort")
    ordered = rows["period_id"].astype(str).tolist()
    # If all dates are missing, keep the user's selection order.
    if rows["_sort_date"].isna().all():
        return ids
    # Preserve any ids that were not present in metadata.
    for pid in ids:
        if pid not in ordered:
            ordered.append(pid)
    return ordered


def _metric_delta(current: float, previous: float) -> str:
    try:
        current = float(current or 0)
        previous = float(previous or 0)
    except Exception:
        return "0"
    diff = current - previous
    sign = "+" if diff > 0 else ""
    if previous:
        pct = diff / previous * 100
        pct_sign = "+" if pct > 0 else ""
        return f"{sign}{format_int(diff)} ({pct_sign}{pct:.0f}%)"
    if diff:
        return f"{sign}{format_int(diff)}"
    return "0"


def _pp_delta(current_share: float, previous_share: float) -> str:
    diff = (float(current_share or 0) - float(previous_share or 0)) * 100
    sign = "+" if diff > 0 else ""
    return f"{sign}{diff:.1f} п.п."


def _period_metrics_for_comparison(messages: pd.DataFrame, periods: pd.DataFrame, period_ids: list[str]) -> list[dict[str, Any]]:
    ordered_ids = _ordered_period_ids(periods, period_ids)
    result: list[dict[str, Any]] = []
    if len(ordered_ids) < 2:
        return result

    rows = _selected_period_rows(periods, ordered_ids)
    row_by_id = {str(r.get("period_id")): r for _, r in rows.iterrows()} if not rows.empty else {}
    for pid in ordered_ids:
        if isinstance(messages, pd.DataFrame) and not messages.empty and "period_id" in messages.columns:
            subset = messages[messages["period_id"].astype(str) == str(pid)].copy()
        else:
            subset = pd.DataFrame()
        metrics = overview_metrics(subset)
        sent = metrics.get("sentiment", {})
        total = max(1, int(sent.get("total", 0) or 0))
        row = row_by_id.get(str(pid), pd.Series({"period_id": pid}))
        metrics.update({
            "period_id": str(pid),
            "label": _period_row_label(row, str(pid)),
            "positive_share": float(sent.get("positive", 0) or 0) / total if total else 0.0,
            "neutral_share": float(sent.get("neutral", 0) or 0) / total if total else 0.0,
            "negative_share": float(sent.get("negative", 0) or 0) / total if total else 0.0,
        })
        result.append(metrics)
    return result


def _comparison_row(metric: dict[str, Any], previous: dict[str, Any] | None = None) -> dict[str, Any]:
    sent = metric.get("sentiment", {}) or {}
    row = {
        "Период": metric.get("label", metric.get("period_id", "")),
        "Сообщений": format_int(metric.get("messages", 0)),
        "Δ сообщений": "—" if previous is None else _metric_delta(metric.get("messages", 0), previous.get("messages", 0)),
        "Аудитория": format_int(metric.get("audience", 0)),
        "Δ аудитории": "—" if previous is None else _metric_delta(metric.get("audience", 0), previous.get("audience", 0)),
        "Охват": format_int(metric.get("reach", 0)),
        "Δ охвата": "—" if previous is None else _metric_delta(metric.get("reach", 0), previous.get("reach", 0)),
        "Вовлеченность": format_int(metric.get("engagement", 0)),
        "Δ вовлеченности": "—" if previous is None else _metric_delta(metric.get("engagement", 0), previous.get("engagement", 0)),
        "Позитив": percent_text(sent.get("positive", 0), sent.get("total", 0)),
        "Δ позитива": "—" if previous is None else _pp_delta(metric.get("positive_share", 0), previous.get("positive_share", 0)),
        "Нейтрал": percent_text(sent.get("neutral", 0), sent.get("total", 0)),
        "Δ нейтрала": "—" if previous is None else _pp_delta(metric.get("neutral_share", 0), previous.get("neutral_share", 0)),
        "Негатив": percent_text(sent.get("negative", 0), sent.get("total", 0)),
        "Δ негатива": "—" if previous is None else _pp_delta(metric.get("negative_share", 0), previous.get("negative_share", 0)),
    }
    return row


def _comparison_visual_rows(comparison: list[dict[str, Any]]) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    for item in comparison:
        sent = item.get("sentiment", {}) or {}
        total = max(1, int(sent.get("total", 0) or 0))
        rows.append({
            "Период": item.get("label", item.get("period_id", "")),
            "Сообщения": int(item.get("messages", 0) or 0),
            "Аудитория": int(item.get("audience", 0) or 0),
            "Охват": int(item.get("reach", 0) or 0),
            "Вовлеченность": int(item.get("engagement", 0) or 0),
            "Позитив, %": round(float(sent.get("positive", 0) or 0) / total * 100, 1),
            "Нейтрал, %": round(float(sent.get("neutral", 0) or 0) / total * 100, 1),
            "Негатив, %": round(float(sent.get("negative", 0) or 0) / total * 100, 1),
            "Позитив": int(sent.get("positive", 0) or 0),
            "Нейтрал": int(sent.get("neutral", 0) or 0),
            "Негатив": int(sent.get("negative", 0) or 0),
        })
    return pd.DataFrame(rows)


def _chart_number_label(value: Any, *, percent: bool = False) -> str:
    try:
        number = float(value or 0)
    except Exception:
        number = 0.0
    if percent:
        raw = f"{number:.1f}".replace(".", ",")
        raw = raw[:-2] if raw.endswith(",0") else raw
        return f"{raw}%"
    return format_int(number)


def _render_sentiment_donut(period_label: str, sentiment: dict[str, Any], *, key: str, label_settings: dict[str, Any] | None = None) -> None:
    total = int((sentiment or {}).get("total", 0) or 0)
    if total <= 0:
        st.caption(f"{period_label}: нет данных для круговой диаграммы")
        return
    pie = pd.DataFrame([
        {"Тональность": "Позитив", "Сообщений": int((sentiment or {}).get("positive", 0) or 0)},
        {"Тональность": "Нейтрал", "Сообщений": int((sentiment or {}).get("neutral", 0) or 0)},
        {"Тональность": "Негатив", "Сообщений": int((sentiment or {}).get("negative", 0) or 0)},
    ])
    pie = pie[pie["Сообщений"] > 0]
    if pie.empty:
        st.caption(f"{period_label}: нет данных для круговой диаграммы")
        return
    pie = pie.copy()
    pie["Доля"] = pie["Сообщений"] / total * 100
    pie["Подпись"] = pie.apply(lambda r: f"{r['Тональность']}: {_chart_number_label(r['Доля'], percent=True)}", axis=1)
    color_map = dict(zip(SENTIMENT_COLOR_DOMAIN, SENTIMENT_COLOR_RANGE))
    pie["Цвет"] = pie["Тональность"].map(color_map).fillna("#999999")

    # Чтобы подписи не накладывались друг на друга, для круговой диаграммы показываем
    # сам donut отдельно, а значения — списком рядом с диаграммой.
    # Легенда по умолчанию скрыта, потому что справа уже есть блок значений.
    donut_cfg = chart_label_settings_from_project_settings({"chart_label_settings": label_settings or {}})
    donut_legend = alt.Legend(title="Тональность") if donut_cfg.get("show_donut_legend") else None
    base = alt.Chart(pie).encode(
        theta=alt.Theta(field="Сообщений", type="quantitative"),
        color=alt.Color(
            field="Тональность",
            type="nominal",
            scale=alt.Scale(domain=SENTIMENT_COLOR_DOMAIN, range=SENTIMENT_COLOR_RANGE),
            legend=donut_legend,
        ),
        tooltip=[
            "Тональность",
            alt.Tooltip("Сообщений:Q", format=","),
            alt.Tooltip("Доля:Q", format=".1f", title="Доля, %"),
        ],
    )
    arcs = base.mark_arc(innerRadius=50)

    left, right = st.columns([3, 2])
    with left:
        st.altair_chart(arcs.properties(height=260, title=period_label), use_container_width=True)
    with right:
        st.markdown("**Значения**")
        for _, row in pie.sort_values("Сообщений", ascending=False).iterrows():
            tone = str(row.get("Тональность", ""))
            color = str(row.get("Цвет", "#999999"))
            count = format_int(row.get("Сообщений", 0))
            share = _chart_number_label(row.get("Доля", 0), percent=True)
            st.markdown(
                f"<div style='margin: 0 0 8px 0; line-height:1.35'>"
                f"<span style='color:{color}; font-size:18px;'>●</span> "
                f"<span style='font-weight:600'>{tone}</span><br>"
                f"<span style='color:#666'>Сообщений:</span> {count}<br>"
                f"<span style='color:#666'>Доля:</span> {share}"
                f"</div>",
                unsafe_allow_html=True,
            )


def _render_value_distribution_donut(df: pd.DataFrame, label_col: str, value_col: str, title: str) -> None:
    if df is None or df.empty or label_col not in df.columns or value_col not in df.columns:
        st.info("Нет данных для круговой диаграммы.")
        return
    pie_df = df[[label_col, value_col]].copy()
    pie_df[value_col] = pd.to_numeric(pie_df[value_col], errors="coerce").fillna(0)
    pie_df = pie_df[pie_df[value_col] > 0]
    total_value = float(pie_df[value_col].sum() or 0)
    if pie_df.empty or total_value <= 0:
        st.info("Нет данных для круговой диаграммы.")
        return
    pie_df["Доля"] = pie_df[value_col] / total_value * 100
    donut = alt.Chart(pie_df).mark_arc(innerRadius=50).encode(
        theta=alt.Theta(field=value_col, type="quantitative"),
        color=alt.Color(f"{label_col}:N", legend=None),
        tooltip=[label_col, alt.Tooltip(f"{value_col}:Q", format=","), alt.Tooltip("Доля:Q", format=".1f", title="Доля, %")],
    )
    left, right = st.columns([3, 2])
    with left:
        st.altair_chart(donut.properties(height=300, title=title), use_container_width=True)
    with right:
        st.markdown("**Значения**")
        for _, row in pie_df.sort_values(value_col, ascending=False).iterrows():
            st.markdown(
                f"**{row[label_col]}**  \n"
                f"{_chart_number_label(row[value_col])} · {_chart_number_label(row['Доля'], percent=True)}"
            )


def render_period_comparison_charts(comparison: list[dict[str, Any]], *, label_settings: dict[str, Any] | None = None) -> None:
    if not comparison:
        return
    chart_df = _comparison_visual_rows(comparison)
    if chart_df.empty:
        return

    st.subheader("Визуализация сравнений")
    st.caption("Выберите, какие графики показать на стартовой странице. Скрытые графики не рендерятся и не перегружают страницу.")

    chart_blocks = [
        "Динамика основных метрик",
        "Динамика тональности",
        "Сравнение выбранной метрики",
        "Круговые диаграммы тональности",
    ]
    selected_blocks = st.multiselect(
        "Показывать графики",
        chart_blocks,
        default=["Динамика основных метрик", "Динамика тональности"],
        key=f"comparison_visible_charts_{abs(hash(tuple(chart_df['Период'].astype(str).tolist())))}",
        help="Можно оставить только нужные визуализации. Это ускоряет отображение страницы при большом числе периодов.",
    )

    if not selected_blocks:
        st.info("Все графики скрыты. Выберите хотя бы один график в списке выше.")
        return

    metrics_cols = ["Сообщения", "Аудитория", "Охват", "Вовлеченность"]

    if "Динамика основных метрик" in selected_blocks:
        st.markdown("**Динамика основных метрик**")
        metrics_long = chart_df[["Период"] + metrics_cols].melt(
            id_vars="Период",
            var_name="Метрика",
            value_name="Значение",
        )
        metrics_long["Подпись"] = metrics_long["Значение"].apply(_chart_number_label)
        chart_type = st.selectbox(
            "Тип визуализации основных метрик",
            ["График", "Столбчатая", "Круговая диаграмма"],
            index=0,
            key=f"main_metrics_chart_type_{abs(hash(tuple(chart_df['Период'].tolist())))}",
        )
        base_metrics = alt.Chart(metrics_long).encode(
            x=alt.X("Период:N", sort=None, title="Период"),
            y=alt.Y("Значение:Q", title="Значение"),
            color=alt.Color("Метрика:N", legend=alt.Legend(title="Метрика")),
            tooltip=["Период", "Метрика", alt.Tooltip("Значение:Q", format=",")],
        )
        if chart_type == "Столбчатая":
            bars = alt.Chart(metrics_long).mark_bar(size=18).encode(
                x=alt.X("Период:N", sort=None, title="Период", axis=alt.Axis(labelAngle=-90)),
                xOffset=alt.XOffset("Метрика:N"),
                y=alt.Y("Значение:Q", title="Значение"),
                color=alt.Color("Метрика:N", legend=alt.Legend(title="Метрика")),
                tooltip=["Период", "Метрика", alt.Tooltip("Значение:Q", format=",")],
            )
            st.altair_chart(bars.properties(height=320), use_container_width=True)
        elif chart_type == "Круговая диаграмма":
            pie_metric = st.selectbox(
                "Метрика для круговой диаграммы",
                metrics_cols,
                index=0,
                key=f"main_metrics_pie_metric_{abs(hash(tuple(chart_df['Период'].tolist())))}",
            )
            pie_df = chart_df[["Период", pie_metric]].rename(columns={pie_metric: "Значение"}).copy()
            pie_df["Значение"] = pd.to_numeric(pie_df["Значение"], errors="coerce").fillna(0)
            total_value = float(pie_df["Значение"].sum() or 0)
            if total_value <= 0:
                st.info("Нет данных для круговой диаграммы по выбранной метрике.")
            else:
                pie_df["Доля"] = pie_df["Значение"] / total_value * 100
                pie_df["Подпись"] = pie_df.apply(lambda r: f"{_chart_number_label(r['Значение'])} · {_chart_number_label(r['Доля'], percent=True)}", axis=1)
                donut = alt.Chart(pie_df).mark_arc(innerRadius=50).encode(
                    theta=alt.Theta(field="Значение", type="quantitative"),
                    color=alt.Color("Период:N", legend=None),
                    tooltip=["Период", alt.Tooltip("Значение:Q", format=","), alt.Tooltip("Доля:Q", format=".1f", title="Доля, %")],
                )
                left_pie, right_pie = st.columns([3, 2])
                with left_pie:
                    st.altair_chart(donut.properties(height=300, title=pie_metric), use_container_width=True)
                with right_pie:
                    st.markdown("**Значения**")
                    for _, row in pie_df.sort_values("Значение", ascending=False).iterrows():
                        st.markdown(f"**{row['Период']}**  \n{_chart_number_label(row['Значение'])} · {_chart_number_label(row['Доля'], percent=True)}")
        else:
            metrics_line = base_metrics.mark_line(point=True)
            metrics_labels = base_metrics.mark_text(**chart_label_text_kwargs(label_settings, chart_type="line")).encode(text="Подпись:N")
            st.altair_chart((metrics_line + metrics_labels).properties(height=320), use_container_width=True)

    if "Динамика тональности" in selected_blocks:
        st.markdown("**Динамика долей тональности, %**")
        sentiment_long = chart_df[["Период", "Позитив, %", "Нейтрал, %", "Негатив, %"]].melt(
            id_vars="Период",
            var_name="Тональность",
            value_name="Доля, %",
        )
        sentiment_long["Тональность"] = sentiment_long["Тональность"].str.replace(", %", "", regex=False)
        sentiment_long["Подпись"] = sentiment_long["Доля, %"].apply(lambda x: _chart_number_label(x, percent=True))
        sentiment_chart_type = st.selectbox(
            "Тип визуализации тональности",
            ["График", "Столбчатая", "Круговая диаграмма"],
            index=0,
            key=f"sentiment_chart_type_{abs(hash(tuple(chart_df['Период'].tolist())))}",
        )
        base_sentiment = alt.Chart(sentiment_long).encode(
            x=alt.X("Период:N", sort=None, title="Период"),
            y=alt.Y("Доля, %:Q", title="Доля, %"),
            color=alt.Color(
                "Тональность:N",
                scale=alt.Scale(domain=SENTIMENT_COLOR_DOMAIN, range=SENTIMENT_COLOR_RANGE),
                legend=alt.Legend(title="Тональность"),
            ),
            tooltip=["Период", "Тональность", alt.Tooltip("Доля, %:Q", format=".1f")],
        )
        if sentiment_chart_type == "Столбчатая":
            sentiment_bars = alt.Chart(sentiment_long).mark_bar(size=18).encode(
                x=alt.X("Период:N", sort=None, title="Период", axis=alt.Axis(labelAngle=-90)),
                xOffset=alt.XOffset("Тональность:N"),
                y=alt.Y("Доля, %:Q", title="Доля, %"),
                color=alt.Color(
                    "Тональность:N",
                    scale=alt.Scale(domain=SENTIMENT_COLOR_DOMAIN, range=SENTIMENT_COLOR_RANGE),
                    legend=alt.Legend(title="Тональность"),
                ),
                tooltip=["Период", "Тональность", alt.Tooltip("Доля, %:Q", format=".1f")],
            )
            st.altair_chart(sentiment_bars.properties(height=320), use_container_width=True)
        elif sentiment_chart_type == "Круговая диаграмма":
            period_options = [str(x) for x in chart_df["Период"].tolist()]
            selected_period_for_sentiment = st.selectbox(
                "Период для круговой диаграммы тональности",
                period_options,
                index=len(period_options) - 1 if period_options else 0,
                key=f"sentiment_pie_period_{abs(hash(tuple(period_options)))}",
            )
            sentiment_by_label = {str(item.get("label", item.get("period_id", ""))): item.get("sentiment", {}) for item in comparison}
            _render_sentiment_donut(selected_period_for_sentiment, sentiment_by_label.get(selected_period_for_sentiment, {}), key="sentiment_selector", label_settings=label_settings)
        else:
            sentiment_line = base_sentiment.mark_line(point=True)
            sentiment_labels = base_sentiment.mark_text(**chart_label_text_kwargs(label_settings, chart_type="line")).encode(text="Подпись:N")
            st.altair_chart((sentiment_line + sentiment_labels).properties(height=320), use_container_width=True)

    if "Сравнение выбранной метрики" in selected_blocks:
        st.markdown("**Сравнение выбранной метрики по периодам**")
        metric_map = {
            "Сообщения": "Сообщения",
            "Аудитория": "Аудитория",
            "Охват": "Охват",
            "Вовлеченность": "Вовлеченность",
        }
        selected_metric = st.selectbox(
            "Метрика для сравнения",
            list(metric_map.keys()),
            index=0,
            key=f"comparison_metric_{abs(hash(tuple(chart_df['Период'].tolist())))}",
        )
        metric_col = metric_map[selected_metric]
        bar_df = chart_df[["Период", metric_col]].rename(columns={metric_col: "Значение"})
        bar_df["Подпись"] = bar_df["Значение"].apply(_chart_number_label)
        comparison_chart_type = st.selectbox(
            "Тип визуализации выбранной метрики",
            ["Столбчатая", "График", "Круговая диаграмма"],
            index=0,
            key=f"single_metric_chart_type_{abs(hash(tuple(chart_df['Период'].tolist())))}",
        )
        if comparison_chart_type == "Круговая диаграмма":
            _render_value_distribution_donut(bar_df, "Период", "Значение", selected_metric)
        else:
            label_cfg = chart_label_settings_from_project_settings({"chart_label_settings": label_settings or {}})
            if label_cfg.get("position") == "center":
                bar_df["_label_y"] = pd.to_numeric(bar_df["Значение"], errors="coerce").fillna(0) / 2
            elif label_cfg.get("position") == "bottom":
                max_value = float(pd.to_numeric(bar_df["Значение"], errors="coerce").fillna(0).max() or 0)
                bar_df["_label_y"] = max_value * 0.03
            else:
                bar_df["_label_y"] = pd.to_numeric(bar_df["Значение"], errors="coerce").fillna(0)
            bar_base = alt.Chart(bar_df).encode(
                x=alt.X("Период:N", sort=None, title="Период", axis=alt.Axis(labelAngle=-90)),
                y=alt.Y("Значение:Q", title=selected_metric),
                tooltip=["Период", alt.Tooltip("Значение:Q", format=",")],
            )
            if comparison_chart_type == "График":
                line = bar_base.mark_line(point=True)
                line_labels = bar_base.mark_text(**chart_label_text_kwargs(label_settings, chart_type="line")).encode(text="Подпись:N")
                st.altair_chart((line + line_labels).properties(height=320), use_container_width=True)
            else:
                bar = bar_base.mark_bar(size=70)
                bar_label_kwargs = chart_label_text_kwargs(label_settings, chart_type="bar")
                if label_cfg.get("position") in {"center", "bottom"}:
                    bar_label_kwargs["dy"] = 0
                    bar_label_kwargs["baseline"] = "middle" if label_cfg.get("position") == "center" else "bottom"
                bar_labels = alt.Chart(bar_df).mark_text(**bar_label_kwargs).encode(
                    x=alt.X("Период:N", sort=None),
                    y=alt.Y("_label_y:Q"),
                    text="Подпись:N",
                )
                st.altair_chart((bar + bar_labels).properties(height=320), use_container_width=True)

    if "Круговые диаграммы тональности" in selected_blocks and len(comparison) >= 2:
        st.markdown("**Круговые диаграммы тональности**")
        unique_periods: list[dict[str, Any]] = []
        seen_periods: set[str] = set()
        for item in comparison:
            key = str(item.get("period_id") or item.get("label") or "").strip()
            if not key:
                key = str(len(unique_periods))
            if key in seen_periods:
                continue
            seen_periods.add(key)
            unique_periods.append(item)

        for row_start in range(0, len(unique_periods), 2):
            cols = st.columns(2)
            for offset, item in enumerate(unique_periods[row_start:row_start + 2]):
                with cols[offset]:
                    donut_key = str(item.get("period_id") or item.get("label") or f"period_{row_start + offset}")
                    _render_sentiment_donut(
                        str(item.get("label", "Период")),
                        item.get("sentiment", {}),
                        key=f"sentiment_donut_{abs(hash(donut_key))}",
                        label_settings=label_settings,
                    )



def render_period_comparison_metrics(messages: pd.DataFrame, periods: pd.DataFrame, period_ids: list[str], *, chart_label_settings: dict[str, Any] | None = None) -> dict[str, Any] | None:
    """Render sequential comparison when two or more periods are selected."""
    comparison = _period_metrics_for_comparison(messages, periods, period_ids)
    if len(comparison) < 2:
        return None

    previous, current = comparison[-2], comparison[-1]
    first, last = comparison[0], comparison[-1]
    st.subheader("Последовательное сравнение периодов")
    st.caption(
        "Сравнение идет цепочкой по хронологии: "
        + " → ".join(item.get("label", item.get("period_id", "")) for item in comparison)
    )

    st.markdown(f"**Последний период:** {current['label']} · сравнение с предыдущим: {previous['label']}")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Сообщений", format_int(current["messages"]), delta=_metric_delta(current["messages"], previous["messages"]))
    c2.metric("Аудитория", format_int(current["audience"]), delta=_metric_delta(current["audience"], previous["audience"]))
    c3.metric("Охват", format_int(current["reach"]), delta=_metric_delta(current["reach"], previous["reach"]))
    c4.metric("Вовлеченность", format_int(current["engagement"]), delta=_metric_delta(current["engagement"], previous["engagement"]))

    s1, s2, s3 = st.columns(3)
    s1.metric(
        "Позитив",
        f"{current['positive_share'] * 100:.0f}%",
        delta=_pp_delta(current["positive_share"], previous["positive_share"]),
        help=f"{format_int(current['sentiment'].get('positive', 0))} сообщений в последнем периоде",
    )
    s2.metric(
        "Нейтрал",
        f"{current['neutral_share'] * 100:.0f}%",
        delta=_pp_delta(current["neutral_share"], previous["neutral_share"]),
        help=f"{format_int(current['sentiment'].get('neutral', 0))} сообщений в последнем периоде",
    )
    s3.metric(
        "Негатив",
        f"{current['negative_share'] * 100:.0f}%",
        delta=_pp_delta(current["negative_share"], previous["negative_share"]),
        help=f"{format_int(current['sentiment'].get('negative', 0))} сообщений в последнем периоде",
    )

    render_period_comparison_charts(comparison, label_settings=chart_label_settings)

    table_rows = []
    prev_item: dict[str, Any] | None = None
    for item in comparison:
        table_rows.append(_comparison_row(item, prev_item))
        prev_item = item

    with st.expander("Последовательная сравнительная таблица", expanded=True):
        st.dataframe(pd.DataFrame(table_rows), hide_index=True, use_container_width=True)

    if len(comparison) > 2:
        st.caption(
            f"Итоговая динамика от первого к последнему периоду: "
            f"сообщения — {_metric_delta(last['messages'], first['messages'])}; "
            f"аудитория — {_metric_delta(last['audience'], first['audience'])}; "
            f"охват — {_metric_delta(last['reach'], first['reach'])}; "
            f"вовлеченность — {_metric_delta(last['engagement'], first['engagement'])}."
        )

    aggregate_metrics = overview_metrics(messages)
    aggregate_metrics["period_label"] = selected_period_label(periods, period_ids)
    aggregate_metrics["comparison_sequence"] = comparison
    aggregate_metrics["comparison"] = {"first": first, "previous": previous, "current": current, "last": last}
    return aggregate_metrics


def render_project_intro(project_name: str, messages: pd.DataFrame, periods: pd.DataFrame, period_ids: list[str], *, profile_label: str = "", chart_label_settings: dict[str, Any] | None = None) -> dict[str, Any]:
    """Unified top block for all project profiles.

    If several periods are selected, the top cards show aggregate values for
    the whole selected range. Sequential comparison is rendered below as a
    separate analytical block and does not replace the aggregate overview.
    """
    st.header(project_name)
    if profile_label:
        st.caption(f"Профиль проекта: {profile_label}")
    period_label = selected_period_label(periods, period_ids)
    selected_ids = [x for x in (period_ids or []) if str(x).strip()]

    st.subheader("Период и основные метрики")
    metrics = overview_metrics(messages)
    sent = metrics["sentiment"]
    total = int(sent.get("total", 0))

    if len(selected_ids) >= 2:
        st.caption(f"Выбрано периодов: {len(selected_ids)} · общие данные по выбранным периодам: {period_label}")
    else:
        st.caption(f"Период: {period_label}")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Сообщений", format_int(metrics["messages"]))
    c2.metric("Суммарная аудитория", format_int(metrics["audience"]))
    c3.metric("Суммарный охват", format_int(metrics["reach"]))
    c4.metric("Суммарная вовлеченность", format_int(metrics["engagement"]))

    t1, t2, t3 = st.columns(3)
    t1.metric("Позитив", f"{percent_text(sent.get('positive', 0), total)}", help=f"{format_int(sent.get('positive', 0))} сообщений")
    t2.metric("Нейтрал", f"{percent_text(sent.get('neutral', 0), total)}", help=f"{format_int(sent.get('neutral', 0))} сообщений")
    t3.metric("Негатив", f"{percent_text(sent.get('negative', 0), total)}", help=f"{format_int(sent.get('negative', 0))} сообщений")

    metrics["period_label"] = period_label
    metrics["project_name"] = project_name

    if len(selected_ids) >= 2 and isinstance(messages, pd.DataFrame) and "period_id" in messages.columns:
        st.divider()
        comparison_metrics = render_period_comparison_metrics(messages, periods, period_ids, chart_label_settings=chart_label_settings)
        if comparison_metrics is not None:
            metrics["comparison_sequence"] = comparison_metrics.get("comparison_sequence")
            metrics["comparison"] = comparison_metrics.get("comparison")

    return metrics


def clean_summary_for_export(text: str) -> str:
    value = str(text or "").strip()
    value = value.replace("**", "")
    value = re.sub(r"^\s*•\s*", "", value, flags=re.MULTILINE)
    return value


def safe_export_filename(project_name: str, period_label: str, ext: str) -> str:
    raw = f"summary_{project_name}_{period_label}"
    safe = re.sub(r"[^0-9A-Za-zА-Яа-я_.-]+", "_", raw, flags=re.UNICODE).strip("_")
    return f"{safe[:140] or 'summary'}.{ext}"


def export_top_tags(messages: pd.DataFrame | None, limit: int = 5) -> list[dict[str, Any]]:
    if messages is None or not isinstance(messages, pd.DataFrame) or messages.empty:
        return []
    try:
        stats = build_tag_statistics(messages).head(limit).copy()
    except Exception:
        return []
    result: list[dict[str, Any]] = []
    for _, row in stats.iterrows():
        tag = str(row.get("Тег") or "").strip()
        if not tag:
            continue
        result.append({
            "name": tag,
            "messages": int(row.get("Сообщений", 0) or 0),
            "audience": int(row.get("Аудитория", 0) or 0),
            "reach": int(row.get("Охват", 0) or 0),
            "engagement": int(row.get("Вовлеченность", 0) or 0),
        })
    return result


def export_top_events(events: pd.DataFrame | None, limit: int = 5) -> list[dict[str, Any]]:
    if events is None or not isinstance(events, pd.DataFrame) or events.empty:
        return []
    work = events.copy()
    title_col = first_existing_col(work, ["display_title", "event_title", "title", "Сюжет / инфоповод", "Сюжет"])
    if title_col is None:
        return []
    work["_title"] = work[title_col].fillna("").astype(str).str.strip()
    # Технические категории не должны попадать в клиентскую инфографику.
    work = work[(work["_title"] != "") & (~work["_title"].str.lower().isin({"без сюжета", "без_сюжета", "без темы", "прочее"}))]
    if work.empty:
        return []
    count_col = first_existing_col(work, ["message_count", "messages", "Сообщений"])
    reach_col = first_existing_col(work, ["views", "reach", "Охват", "Просмотры", "Просмотров"])
    engagement_col = first_existing_col(work, ["engagement", "Вовлеченность", "Вовлечённость"])
    work["_messages"] = numeric_series(work, [count_col]).astype(int) if count_col else 0
    work["_reach"] = numeric_series(work, [reach_col]).astype(int) if reach_col else 0
    work["_engagement"] = numeric_series(work, [engagement_col]).astype(int) if engagement_col else 0
    work = work.sort_values(["_messages", "_reach", "_engagement"], ascending=False).head(limit)
    result: list[dict[str, Any]] = []
    for _, row in work.iterrows():
        result.append({
            "name": str(row.get("_title") or "").strip(),
            "messages": int(row.get("_messages", 0) or 0),
            "reach": int(row.get("_reach", 0) or 0),
            "engagement": int(row.get("_engagement", 0) or 0),
        })
    return result


def summary_highlights(summary_text: str, limit: int = 4) -> list[str]:
    lines: list[str] = []
    for raw in str(summary_text or "").replace("\r", "\n").split("\n"):
        line = raw.strip().strip("•-").strip()
        if line and line not in lines:
            lines.append(line)
        if len(lines) >= limit:
            break
    if lines:
        return lines
    return ["Саммари пока не заполнено."]


def summary_export_payload(
    project_name: str,
    period_label: str,
    summary_text: str,
    metrics: dict[str, Any],
    messages: pd.DataFrame | None = None,
    events_agg: pd.DataFrame | None = None,
    *,
    report_template: str = "summary",
    branding: dict[str, Any] | None = None,
) -> dict[str, Any]:
    sent = metrics.get("sentiment", {}) if isinstance(metrics, dict) else {}
    report_template = report_template if report_template in REPORT_TEMPLATE_OPTIONS else "summary"
    branding = report_branding_from_project_settings({"report_branding": branding or {}}, project_name=project_name)
    return {
        "project_name": project_name,
        "client_name": branding.get("client_name") or project_name,
        "report_title": branding.get("report_title") or "Дайджест упоминаний",
        "accent_color": branding.get("accent_color") or "#2563eb",
        "background_color": branding.get("background_color") or "#ffffff",
        "footer_text": branding.get("footer_text") or "",
        "logo_url": branding.get("logo_url") or "",
        "report_template": report_template,
        "report_template_label": REPORT_TEMPLATE_OPTIONS.get(report_template, report_template),
        "period_label": period_label,
        "summary_text": clean_summary_for_export(summary_text),
        "summary_highlights": summary_highlights(summary_text),
        "messages": int(metrics.get("messages", 0) or 0),
        "audience": int(metrics.get("audience", 0) or 0),
        "reach": int(metrics.get("reach", 0) or 0),
        "engagement": int(metrics.get("engagement", 0) or 0),
        "positive": int(sent.get("positive", 0) or 0),
        "neutral": int(sent.get("neutral", 0) or 0),
        "negative": int(sent.get("negative", 0) or 0),
        "total": int(sent.get("total", 0) or 0),
        "comparison_sequence": metrics.get("comparison_sequence") or [],
        "top_tags": export_top_tags(messages, limit=8 if report_template == "full" else 5),
        "top_events": export_top_events(events_agg, limit=8 if report_template == "full" else 5),
        "created_at": datetime.now().strftime("%d.%m.%Y %H:%M"),
    }


def _short_label(value: Any, max_len: int = 34) -> str:
    text = str(value or "").strip()
    return text if len(text) <= max_len else text[: max_len - 1].rstrip() + "…"


def _metric_delta_for_export(current: Any, previous: Any) -> str:
    try:
        cur = float(current or 0)
        prev = float(previous or 0)
    except Exception:
        return ""
    diff = cur - prev
    if prev:
        pct = diff / abs(prev) * 100
        return f"{diff:+,.0f} / {pct:+.1f}%".replace(",", " ")
    if diff:
        return f"{diff:+,.0f}".replace(",", " ")
    return "0"


def _draw_export_card(ax, x: float, y: float, w: float, h: float, title: str, value: str, subtitle: str = "", *, accent_color: str = "#2563eb") -> None:
    import matplotlib.pyplot as plt  # noqa: F401
    from matplotlib.patches import FancyBboxPatch
    patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.012,rounding_size=0.018", linewidth=1, edgecolor="#d9e0ea", facecolor="#f8fafc")
    ax.add_patch(patch)
    ax.plot([x + 0.015, x + 0.015], [y + 0.018, y + h - 0.018], color=accent_color, linewidth=2.4)
    ax.text(x + 0.032, y + h - 0.035, title, fontsize=9.5, color="#4b5563", va="top", ha="left")
    ax.text(x + 0.032, y + h / 2 + 0.005, value, fontsize=15, color="#111827", va="center", ha="left", fontweight="bold")
    if subtitle:
        ax.text(x + 0.032, y + 0.018, subtitle, fontsize=8.4, color="#6b7280", va="bottom", ha="left")


def generate_summary_infographic_png(payload: dict[str, Any]) -> bytes:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import Rectangle
    except Exception as exc:
        raise RuntimeError("Для инфографики нужен matplotlib>=3.8 в requirements.txt.") from exc

    plt.rcParams["font.family"] = "DejaVu Sans"
    fig = plt.figure(figsize=(8.27, 11.69), dpi=170)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    background = _valid_hex_color(payload.get("background_color"), "#ffffff")
    accent = _valid_hex_color(payload.get("accent_color"), "#2563eb")
    fig.patch.set_facecolor(background)

    project = _short_label(payload.get("client_name") or payload.get("project_name") or "Проект", 52)
    report_title = _short_label(payload.get("report_title") or "Дайджест упоминаний", 52)
    template_label = str(payload.get("report_template_label") or "")
    period = _short_label(payload.get("period_label") or "выбранный период", 70)
    created = str(payload.get("created_at") or "")
    comparison = payload.get("comparison_sequence") or []

    ax.add_patch(Rectangle((0, 0.895), 1, 0.105, facecolor=accent, edgecolor="none", alpha=0.95))
    ax.text(0.06, 0.972, report_title, fontsize=13, fontweight="bold", color="white", va="top", ha="left")
    ax.text(0.06, 0.945, project, fontsize=21, fontweight="bold", color="white", va="top", ha="left")
    ax.text(0.06, 0.918, f"{template_label} · {period}", fontsize=9.4, color="#e5e7eb", va="top", ha="left")
    ax.text(0.78, 0.972, created, fontsize=8.8, color="#e5e7eb", va="top", ha="left")

    # 2×2-сетка метрик: большие числа больше не накладываются.
    if len(comparison) >= 2:
        previous, current = comparison[-2], comparison[-1]
        metric_cards = [
            ("Сообщения", current.get("messages", 0), _metric_delta_for_export(current.get("messages", 0), previous.get("messages", 0))),
            ("Аудитория", current.get("audience", 0), _metric_delta_for_export(current.get("audience", 0), previous.get("audience", 0))),
            ("Охват", current.get("reach", 0), _metric_delta_for_export(current.get("reach", 0), previous.get("reach", 0))),
            ("Вовлеченность", current.get("engagement", 0), _metric_delta_for_export(current.get("engagement", 0), previous.get("engagement", 0))),
        ]
        ax.text(0.06, 0.868, f"Последний период: {_short_label(current.get('label'), 42)}", fontsize=9.2, color="#6b7280", va="top")
    else:
        metric_cards = [
            ("Сообщения", payload.get("messages", 0), ""),
            ("Аудитория", payload.get("audience", 0), ""),
            ("Охват", payload.get("reach", 0), ""),
            ("Вовлеченность", payload.get("engagement", 0), ""),
        ]

    xs = [0.06, 0.52]
    ys = [0.775, 0.67]
    for idx, (title, value, subtitle) in enumerate(metric_cards):
        _draw_export_card(ax, xs[idx % 2], ys[idx // 2], 0.40, 0.08, title, format_int(value), f"к пред. периоду: {subtitle}" if subtitle else "", accent_color=accent)

    total = max(1, int(payload.get("total", 0) or 0))
    pos = int(payload.get("positive", 0) or 0)
    neu = int(payload.get("neutral", 0) or 0)
    neg = int(payload.get("negative", 0) or 0)
    if len(comparison) >= 2:
        sent = comparison[-1].get("sentiment", {}) or {}
        total = max(1, int(sent.get("total", 0) or 0))
        pos = int(sent.get("positive", 0) or 0)
        neu = int(sent.get("neutral", 0) or 0)
        neg = int(sent.get("negative", 0) or 0)

    ax.text(0.06, 0.605, "Тональность", fontsize=13, fontweight="bold", color="#111827", va="top")
    pie_ax = fig.add_axes([0.06, 0.43, 0.24, 0.16])
    pie_ax.axis("equal")
    values = [max(pos, 0), max(neu, 0), max(neg, 0)]
    colors = ["#22c55e", "#9ca3af", "#ef4444"]
    labels = ["Позитив", "Нейтрал", "Негатив"]
    pie_ax.pie(values, colors=colors, startangle=90, counterclock=False, wedgeprops={"width": 0.42, "edgecolor": "white"})
    pie_ax.text(0, 0.05, format_int(total), ha="center", va="center", fontsize=14, fontweight="bold", color="#111827")
    pie_ax.text(0, -0.12, "сообщений", ha="center", va="center", fontsize=8.5, color="#6b7280")
    pie_ax.set_xticks([])
    pie_ax.set_yticks([])

    y0 = 0.565
    for i, (lab, val, col) in enumerate(zip(labels, values, colors)):
        yy = y0 - i * 0.044
        ax.add_patch(Rectangle((0.33, yy - 0.011), 0.016, 0.016, facecolor=col, edgecolor="none"))
        ax.text(0.355, yy, lab, fontsize=10, color="#111827", va="center", ha="left")
        ax.text(0.50, yy, format_int(val), fontsize=10, color="#111827", va="center", ha="right", fontweight="bold")
        ax.text(0.52, yy, percent_text(val, total), fontsize=9, color="#6b7280", va="center", ha="left")

    # Топы добавляют аналитическую ценность, а не просто дублируют KPI.
    top_tags = payload.get("top_tags") or []
    top_events = payload.get("top_events") or []
    ax.text(0.06, 0.395, "Топ тегов", fontsize=12.5, fontweight="bold", color="#111827", va="top")
    y = 0.367
    if top_tags:
        for item in top_tags[:5]:
            ax.text(0.07, y, f"• {_short_label(item.get('name'), 36)}", fontsize=9.2, color="#111827", va="top")
            ax.text(0.43, y, f"{format_int(item.get('messages', 0))} сообщ.", fontsize=8.8, color="#6b7280", va="top", ha="right")
            y -= 0.03
    else:
        ax.text(0.07, y, "Нет тегов для отображения", fontsize=9.2, color="#6b7280", va="top")

    ax.text(0.52, 0.395, "Топ инфоповодов", fontsize=12.5, fontweight="bold", color="#111827", va="top")
    y = 0.367
    if top_events:
        for item in top_events[:5]:
            ax.text(0.53, y, f"• {_short_label(item.get('name'), 38)}", fontsize=9.2, color="#111827", va="top")
            ax.text(0.92, y, f"{format_int(item.get('messages', 0))} сообщ.", fontsize=8.8, color="#6b7280", va="top", ha="right")
            y -= 0.03
    else:
        ax.text(0.53, y, "Нет инфоповодов для отображения", fontsize=9.2, color="#6b7280", va="top")

    ax.text(0.06, 0.215, "Главное", fontsize=12.5, fontweight="bold", color="#111827", va="top")
    summary_y = 0.188
    line_count = 0
    for block in (payload.get("summary_highlights") or [])[:4]:
        wrapped = textwrap.wrap(str(block), width=90) or [str(block)]
        bullet = True
        for seg in wrapped[:2]:
            prefix = "• " if bullet else "  "
            ax.text(0.07, summary_y, prefix + seg, fontsize=9.4, color="#111827", va="top", ha="left")
            summary_y -= 0.024
            line_count += 1
            bullet = False
            if line_count >= 8:
                break
        summary_y -= 0.005
        if line_count >= 8:
            break

    footer_text = str(payload.get("footer_text") or "Инфографика сформирована автоматически на основе выбранного периода и текущего саммари.")
    ax.text(0.06, 0.045, footer_text, fontsize=8.4, color="#6b7280", va="bottom", ha="left")

    out = BytesIO()
    fig.savefig(out, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out.getvalue()


def generate_summary_docx(payload: dict[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except Exception as exc:
        raise RuntimeError("Для выгрузки Word добавьте python-docx в requirements.txt.") from exc

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    r = p.add_run(str(payload.get("report_title") or "Дайджест упоминаний"))
    r.bold = True
    r.font.size = Pt(16)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(str(payload.get("client_name") or payload.get("project_name") or "Проект"))
    r2.bold = True
    r2.font.size = Pt(13)
    doc.add_paragraph(f"Шаблон: {payload.get('report_template_label') or ''}")
    doc.add_paragraph(f"Период: {payload.get('period_label') or 'выбранный период'}")
    doc.add_paragraph(f"Дата выгрузки: {payload.get('created_at') or ''}")

    total = max(1, int(payload.get("total", 0) or 0))
    doc.add_paragraph(
        "Основные метрики: "
        f"сообщений — {format_int(payload.get('messages', 0))}; "
        f"аудитория — {format_int(payload.get('audience', 0))}; "
        f"охват — {format_int(payload.get('reach', 0))}; "
        f"вовлеченность — {format_int(payload.get('engagement', 0))}."
    )
    doc.add_paragraph(
        "Тональность: "
        f"позитив — {percent_text(int(payload.get('positive', 0) or 0), total)}; "
        f"нейтрал — {percent_text(int(payload.get('neutral', 0) or 0), total)}; "
        f"негатив — {percent_text(int(payload.get('negative', 0) or 0), total)}."
    )

    try:
        infographic_png = generate_summary_infographic_png(payload)
        doc.add_paragraph()
        doc.add_picture(BytesIO(infographic_png), width=Inches(6.2))
    except Exception:
        pass

    if payload.get("report_template") in {"client_overview", "comparison", "full"}:
        p = doc.add_paragraph()
        p.add_run("Что включить в отчет").bold = True
        top_tags = payload.get("top_tags") or []
        top_events = payload.get("top_events") or []
        if top_tags:
            doc.add_paragraph("Топ тегов: " + "; ".join(str(x.get("name") or "") for x in top_tags[:5] if x.get("name")))
        if top_events:
            doc.add_paragraph("Топ инфоповодов: " + "; ".join(str(x.get("name") or "") for x in top_events[:5] if x.get("name")))

    p = doc.add_paragraph()
    p.add_run("Саммари периода").bold = True
    for block in str(payload.get("summary_text") or "").split("\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _pdf_font_candidates() -> list[tuple[str, str | None]]:
    """Return regular/bold TTF candidates that support Cyrillic.

    Streamlit Cloud images do not always include system DejaVu/Noto fonts. If we
    fall back to ReportLab core Helvetica, Cyrillic is rendered as black squares.
    To make PDF export portable, we also look for the DejaVu Sans bundled with
    matplotlib when the package is installed.
    """
    candidates: list[tuple[str, str | None]] = []

    env_path = os.getenv("PLATFORM_PDF_FONT_PATH", "").strip()
    if env_path:
        candidates.append((env_path, None))

    # Optional project-local fonts. Do not commit font files unless licensing is clear;
    # this path is only for private deployments that provide their own font.
    here = Path(__file__).resolve().parent
    candidates.extend([
        (str(here / "assets" / "fonts" / "DejaVuSans.ttf"), str(here / "assets" / "fonts" / "DejaVuSans-Bold.ttf")),
        (str(here.parent / "assets" / "fonts" / "DejaVuSans.ttf"), str(here.parent / "assets" / "fonts" / "DejaVuSans-Bold.ttf")),
    ])

    # Common Linux/Streamlit Cloud system fonts.
    candidates.extend([
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed-Bold.ttf"),
        ("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf", "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/freefont/FreeSans.ttf", "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf"),
        ("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
    ])

    try:
        import matplotlib  # type: ignore

        mpl_fonts = Path(matplotlib.get_data_path()) / "fonts" / "ttf"
        candidates.append((str(mpl_fonts / "DejaVuSans.ttf"), str(mpl_fonts / "DejaVuSans-Bold.ttf")))
    except Exception:
        pass

    return candidates


def _pdf_font_name() -> str:
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as exc:
        raise RuntimeError("Для выгрузки PDF добавьте reportlab в requirements.txt.") from exc

    for regular_path, bold_path in _pdf_font_candidates():
        regular = Path(regular_path)
        if not regular.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("PlatformSans", str(regular)))
            bold = Path(bold_path) if bold_path else None
            if bold is not None and bold.exists():
                pdfmetrics.registerFont(TTFont("PlatformSans-Bold", str(bold)))
                try:
                    pdfmetrics.registerFontFamily(
                        "PlatformSans",
                        normal="PlatformSans",
                        bold="PlatformSans-Bold",
                        italic="PlatformSans",
                        boldItalic="PlatformSans-Bold",
                    )
                except Exception:
                    pass
            return "PlatformSans"
        except Exception:
            continue

    raise RuntimeError(
        "Не найден TTF-шрифт с поддержкой кириллицы для PDF. "
        "Проверьте, что установлен matplotlib>=3.8 или задайте PLATFORM_PDF_FONT_PATH."
    )


def generate_summary_pdf(payload: dict[str, Any]) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
        from xml.sax.saxutils import escape as xml_escape
    except Exception as exc:
        raise RuntimeError("Для выгрузки PDF добавьте reportlab в requirements.txt.") from exc

    font_name = _pdf_font_name()
    out = BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, leftMargin=1.7 * cm, rightMargin=1.7 * cm, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    base = getSampleStyleSheet()
    normal = ParagraphStyle("PlatformNormal", parent=base["Normal"], fontName=font_name, fontSize=10, leading=14)
    title = ParagraphStyle("PlatformTitle", parent=normal, fontName=font_name, fontSize=16, leading=20, spaceAfter=10)
    heading = ParagraphStyle("PlatformHeading", parent=normal, fontName=font_name, fontSize=12, leading=16, spaceBefore=8, spaceAfter=6)

    story = []
    infographic_added = False
    try:
        infographic_png = generate_summary_infographic_png(payload)
        infographic_io = BytesIO(infographic_png)
        infographic_io.seek(0)
        # Инфографика теперь первая страница PDF, без дублирующей текстовой страницы.
        story.append(Image(infographic_io, width=18.0 * cm, height=25.4 * cm))
        story.append(PageBreak())
        infographic_added = True
    except Exception:
        pass

    total = max(1, int(payload.get("total", 0) or 0))
    if not infographic_added:
        story.extend([
            Paragraph(f"<b>{xml_escape(str(payload.get('report_title') or 'Дайджест упоминаний'))}</b>", title),
            Paragraph(xml_escape(str(payload.get('client_name') or payload.get('project_name') or 'Проект')), heading),
            Paragraph(xml_escape(f"Шаблон: {payload.get('report_template_label') or ''}"), normal),
            Paragraph(xml_escape(f"Период: {payload.get('period_label') or 'выбранный период'}"), normal),
            Paragraph(xml_escape(f"Дата выгрузки: {payload.get('created_at') or ''}"), normal),
            Spacer(1, 8),
            Paragraph("<b>Основные метрики</b>", heading),
            Paragraph(xml_escape(
                f"Сообщений — {format_int(payload.get('messages', 0))}; "
                f"аудитория — {format_int(payload.get('audience', 0))}; "
                f"охват — {format_int(payload.get('reach', 0))}; "
                f"вовлеченность — {format_int(payload.get('engagement', 0))}."
            ), normal),
            Paragraph(xml_escape(
                f"Тональность: позитив — {percent_text(int(payload.get('positive', 0) or 0), total)}; "
                f"нейтрал — {percent_text(int(payload.get('neutral', 0) or 0), total)}; "
                f"негатив — {percent_text(int(payload.get('negative', 0) or 0), total)}."
            ), normal),
            Spacer(1, 8),
        ])

    story.append(Paragraph("<b>Саммари периода</b>", heading))
    for block in str(payload.get("summary_text") or "").split("\n"):
        block = block.strip()
        if block:
            story.append(Paragraph(xml_escape(block), normal))
    doc.build(story)
    return out.getvalue()


def render_summary_export_buttons(
    project_name: str,
    period_label: str,
    summary_text: str,
    metrics: dict[str, Any],
    *,
    key_prefix: str,
    messages: pd.DataFrame | None = None,
    events_agg: pd.DataFrame | None = None,
    branding: dict[str, Any] | None = None,
) -> None:
    report_template = st.selectbox(
        "Шаблон отчета",
        list(REPORT_TEMPLATE_OPTIONS.keys()),
        index=0,
        format_func=lambda x: REPORT_TEMPLATE_OPTIONS.get(x, x),
        key=f"{key_prefix}_template",
        help="Шаблон меняет структуру выгрузки и набор аналитических блоков в Word/PDF/PNG.",
    )
    payload = summary_export_payload(
        project_name,
        period_label,
        summary_text,
        metrics,
        messages=messages,
        events_agg=events_agg,
        report_template=report_template,
        branding=branding,
    )
    st.caption(f"Брендирование: {payload.get('client_name') or project_name}; акцентный цвет {payload.get('accent_color')}.")
    c1, c2, c3 = st.columns(3)
    with c1:
        try:
            st.download_button(
                "Скачать Word",
                data=generate_summary_docx(payload),
                file_name=safe_export_filename(project_name, period_label, "docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"{key_prefix}_docx",
            )
        except Exception as exc:
            st.warning(str(exc))
    with c2:
        try:
            st.download_button(
                "Скачать PDF",
                data=generate_summary_pdf(payload),
                file_name=safe_export_filename(project_name, period_label, "pdf"),
                mime="application/pdf",
                use_container_width=True,
                key=f"{key_prefix}_pdf",
            )
        except Exception as exc:
            st.warning(str(exc))
    with c3:
        try:
            st.download_button(
                "Скачать инфографику PNG",
                data=generate_summary_infographic_png(payload),
                file_name=safe_export_filename(project_name, period_label, "png"),
                mime="image/png",
                use_container_width=True,
                key=f"{key_prefix}_png",
            )
        except Exception as exc:
            st.warning(str(exc))


def render_tag_statistics(messages: pd.DataFrame, *, project_id: str | None = None) -> None:
    stats = build_tag_statistics(messages)
    if stats.empty:
        st.info("Теги не найдены.")
        return

    st.subheader("Статистика тегов")
    top = stats.head(30).copy()
    display = top.copy()
    for col in ["Сообщений", "Аудитория", "Охват", "Вовлеченность", "Негатив"]:
        if col in display.columns:
            display[col] = display[col].apply(format_int)
    display["Доля негатива"] = display["Доля негатива"].astype(str) + "%"
    st.caption("Теги берутся из системных колонок Brand Analytics после «Обработано». Аудитория, охват и вовлеченность суммируются по сообщениям с выбранным тегом.")
    tag_selection = st.dataframe(
        display,
        hide_index=True,
        use_container_width=True,
        selection_mode="single-row",
        on_select="rerun",
        key=f"tag_statistics_table_{project_id or 'global'}",
    )
    rows = getattr(tag_selection, "selection", {}).get("rows", []) if tag_selection is not None else []
    if not rows:
        st.caption("Выберите строку тега, чтобы открыть метрики, ключевые сообщения и всю ленту сообщений по этому тегу.")
        return
    selected_tag = str(top.iloc[rows[0]].get("Тег") or "").strip()
    if selected_tag:
        render_selected_tag_detail(project_id or "global", selected_tag, messages)


def messages_with_tag(messages: pd.DataFrame, tag: str) -> pd.DataFrame:
    if messages is None or messages.empty or "tags" not in messages.columns or not str(tag).strip():
        return pd.DataFrame()
    key = str(tag).strip().lower().replace("ё", "е")
    mask = messages["tags"].fillna("").astype(str).apply(
        lambda value: key in {item.lower().replace("ё", "е") for item in split_pipe_values(value)}
    )
    return messages[mask].copy()


def _tag_auto_summary(tag: str, tag_messages: pd.DataFrame) -> str:
    if tag_messages is None or tag_messages.empty:
        return f"По тегу «{tag}» сообщения не найдены."
    metrics = overview_metrics(tag_messages)
    sent = metrics.get("sentiment", {}) or {}
    total = int(sent.get("total", metrics.get("messages", 0)) or 0)
    negative = int(sent.get("negative", 0) or 0)
    event_col = next((col for col in ["event_title", "source_main_topic", "Сюжет"] if col in tag_messages.columns), None)
    event_part = ""
    if event_col:
        top_events = (
            tag_messages[event_col]
            .fillna("")
            .astype(str)
            .str.strip()
            .replace("", pd.NA)
            .dropna()
            .value_counts()
            .head(5)
            .index
            .tolist()
        )
        if top_events:
            event_part = " Основные связанные инфоповоды: " + "; ".join(map(str, top_events)) + "."
    return (
        f"По тегу «{tag}» найдено {format_int(total)} сообщений. "
        f"Негативных сообщений: {format_int(negative)} ({percent_text(negative, total)}). "
        f"Суммарный охват — {format_int(metrics.get('reach', 0))}, "
        f"вовлеченность — {format_int(metrics.get('engagement', 0))}."
        f"{event_part}"
    )


def render_selected_tag_detail(project_id: str, selected_tag: str, messages: pd.DataFrame) -> None:
    """Render selected-tag card with metrics and messages, like selected event card."""
    tag_messages = messages_with_tag(messages, selected_tag)
    st.markdown(f"## Тег: {selected_tag}")
    st.info(_tag_auto_summary(selected_tag, tag_messages))

    if tag_messages is None or tag_messages.empty:
        st.info("Сообщений по выбранному тегу не найдено.")
        return

    metrics = overview_metrics(tag_messages)
    sent = metrics.get("sentiment", {}) or {}
    total = int(sent.get("total", metrics.get("messages", 0)) or 0)

    source_count = 0
    for col in ["chat_title", "platform", "source", "Источник", "Место публикации"]:
        if col in tag_messages.columns:
            source_count = int(tag_messages[col].fillna("").astype(str).replace("", pd.NA).dropna().nunique())
            break
    author_count = 0
    for col in ["author", "Автор"]:
        if col in tag_messages.columns:
            author_count = int(tag_messages[col].fillna("").astype(str).replace("", pd.NA).dropna().nunique())
            break

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Сообщений", format_int(metrics.get("messages", 0)))
    c2.metric("Источников/чатов", format_int(source_count))
    c3.metric("Авторов", format_int(author_count))
    c4.metric("Негатив", percent_text(int(sent.get("negative", 0) or 0), total))

    m1, m2, m3 = st.columns(3)
    m1.metric("Аудитория", format_int(metrics.get("audience", 0)))
    m2.metric("Охват", format_int(metrics.get("reach", 0)))
    m3.metric("Вовлеченность", format_int(metrics.get("engagement", 0)))

    mode = st.radio(
        "Сообщения тега",
        ["Ключевые сообщения", "Вся лента"],
        horizontal=True,
        key=f"selected_tag_messages_mode_{project_id}_{abs(hash(selected_tag))}",
    )

    work = tag_messages.copy()
    text_col = message_text_column(work)
    link_col = message_link_column(work)
    work["_audience"] = numeric_series(work, ["audience", "Аудитория"]).astype(int)
    work["_reach"] = numeric_series(work, ["views", "Просмотры", "Просмотров", "reach", "Охват"]).astype(int)
    work["_engagement"] = numeric_series(work, ["engagement", "Вовлечённость", "Вовлеченность", "engagement_count"]).astype(int)

    if mode == "Ключевые сообщения":
        st.caption("Показаны топ-15 сообщений выбранного тега по вовлеченности. Если вовлеченность равна 0, учитываются охват и аудитория.")
        view = work.sort_values(["_engagement", "_reach", "_audience"], ascending=False).head(15).copy()
    else:
        search_key = f"selected_tag_feed_search_{project_id}_{abs(hash(selected_tag))}"
        search = st.text_input("Поиск по ленте тега", placeholder="Введите слово или фразу", key=search_key)
        view = work.copy()
        if search.strip() and text_col:
            view = view[view[text_col].fillna("").astype(str).str.contains(search.strip(), case=False, regex=False)]
        view = view.sort_values("datetime", ascending=False) if "datetime" in view.columns else view
        total_found = int(len(view))
        page_size = int(st.selectbox(
            "Сообщений на странице",
            [25, 50, 100, 200],
            index=1,
            key=f"selected_tag_feed_page_size_{project_id}_{abs(hash(selected_tag))}",
        ))
        total_pages = max(1, (total_found + page_size - 1) // page_size)
        page = int(st.number_input(
            "Страница",
            min_value=1,
            max_value=total_pages,
            value=1,
            step=1,
            key=f"selected_tag_feed_page_{project_id}_{abs(hash(selected_tag))}",
        ))
        start = (page - 1) * page_size
        end = start + page_size
        st.caption(
            f"Найдено сообщений: {format_int(total_found)}. "
            f"Показано: {format_int(start + 1 if total_found else 0)}–{format_int(min(end, total_found))} из {format_int(total_found)}."
        )
        view = view.iloc[start:end].copy()

    _render_message_list(view, text_col=text_col, link_col=link_col)


def render_tag_explorer(messages: pd.DataFrame, *, key_prefix: str = "tag_explorer") -> None:
    stats = build_tag_statistics(messages)
    if stats.empty:
        return

    st.markdown("#### Теги")
    st.caption("Выберите тег, чтобы посмотреть все сообщения с этим тегом.")
    options = stats["Тег"].astype(str).tolist()
    labels = {
        str(row["Тег"]): f"{row['Тег']} · {format_int(row['Сообщений'])} сообщ. · {format_int(row['Аудитория'])} аудитория · {format_int(row['Охват'])} охват · {format_int(row['Вовлеченность'])} вовлеч."
        for _, row in stats.iterrows()
    }
    selected_tag = st.selectbox(
        "Тег",
        options,
        format_func=lambda x: labels.get(str(x), str(x)),
        key=f"{key_prefix}_select",
    )
    tag_messages = messages_with_tag(messages, selected_tag)
    st.caption(f"Сообщений с тегом «{selected_tag}»: {format_int(len(tag_messages))}")
    if tag_messages.empty:
        return
    tag_messages = tag_messages.copy()
    tag_messages["Дата"] = tag_messages.get("datetime", "").apply(fmt_date)
    tag_messages["Аудитория"] = numeric_series(tag_messages, ["audience", "Аудитория"]).astype(int)
    tag_messages["Охват"] = numeric_series(tag_messages, ["views", "Просмотры", "Просмотров", "reach", "Охват"]).astype(int)
    tag_messages["Вовлеченность"] = numeric_series(tag_messages, ["engagement", "Вовлечённость", "Вовлеченность", "engagement_count"]).astype(int)
    columns = [c for c in ["Дата", "chat_title", "author", "event_title", "text", "url", "Аудитория", "Охват", "Вовлеченность"] if c in tag_messages.columns]
    st.dataframe(
        tag_messages[columns].rename(columns={
            "chat_title": "Источник/площадка",
            "author": "Автор",
            "event_title": "Инфоповод",
            "text": "Текст",
            "url": "Ссылка",
        }),
        hide_index=True,
        use_container_width=True,
        height=420,
        column_config={"Ссылка": st.column_config.LinkColumn("Ссылка")} if "url" in tag_messages.columns else None,
    )



TECHNICAL_EVENT_TITLES = {
    "без сюжета",
    "без темы",
    "прочее",
    "прочие сообщения",
    "общее обсуждение",
}


def is_technical_event_title(title: Any) -> bool:
    value = str(title or "").strip().lower().replace("ё", "е")
    return not value or value in {x.replace("ё", "е") for x in TECHNICAL_EVENT_TITLES}


def _event_title_col(events_agg: pd.DataFrame) -> str | None:
    for col in ["title", "event_title", "Сюжет / инфоповод"]:
        if isinstance(events_agg, pd.DataFrame) and col in events_agg.columns:
            return col
    return None


def top_client_events(events_agg: pd.DataFrame, limit: int = 5) -> pd.DataFrame:
    if events_agg is None or events_agg.empty:
        return pd.DataFrame()
    work = events_agg.copy()
    title_col = _event_title_col(work)
    if not title_col:
        return pd.DataFrame()
    work = work[~work[title_col].apply(is_technical_event_title)].copy()
    if work.empty:
        return work
    for col in ["message_count", "negative_count", "importance_score"]:
        if col in work.columns:
            work[col] = pd.to_numeric(work[col], errors="coerce").fillna(0)
    sort_cols = [c for c in ["message_count", "importance_score"] if c in work.columns]
    if sort_cols:
        work = work.sort_values(sort_cols, ascending=False)
    return work.head(limit).copy()


def top_client_tags(messages: pd.DataFrame, limit: int = 5) -> pd.DataFrame:
    stats = build_tag_statistics(messages)
    if stats is None or stats.empty:
        return pd.DataFrame()
    work = stats.copy()
    for col in ["Сообщений", "Аудитория", "Охват", "Вовлеченность", "Негатив", "Доля негатива"]:
        if col in work.columns:
            work[col] = pd.to_numeric(work[col], errors="coerce").fillna(0)
    sort_cols = [c for c in ["Сообщений", "Охват", "Вовлеченность"] if c in work.columns]
    if sort_cols:
        work = work.sort_values(sort_cols, ascending=False)
    return work.head(limit).copy()


def build_period_change_insights(messages: pd.DataFrame, periods: pd.DataFrame, selected_period_ids: list[str]) -> list[str]:
    comparison = _period_metrics_for_comparison(messages, periods, selected_period_ids)
    if len(comparison) < 2:
        return []
    prev, cur = comparison[-2], comparison[-1]
    insights: list[str] = []
    checks = [
        ("сообщений", "messages"),
        ("аудитории", "audience"),
        ("охвата", "reach"),
        ("вовлеченности", "engagement"),
    ]
    for label, key in checks:
        old = float(prev.get(key, 0) or 0)
        new = float(cur.get(key, 0) or 0)
        if old == 0 and new == 0:
            continue
        delta = new - old
        if abs(delta) < 1:
            continue
        direction = "выросла" if delta > 0 else "снизилась"
        if label == "сообщений":
            direction = "выросло" if delta > 0 else "снизилось"
        percent = f" ({delta / old * 100:+.0f}%)" if old else ""
        insights.append(f"Количество {label} {direction}: {format_int(delta)}{percent} к предыдущему периоду.")

    neg_delta = float(cur.get("negative_share", 0) or 0) - float(prev.get("negative_share", 0) or 0)
    if abs(neg_delta) >= 0.001:
        insights.append(f"Доля негатива изменилась на {neg_delta * 100:+.1f} п.п. к предыдущему периоду.")
    return insights


def build_tag_change_table(messages: pd.DataFrame, periods: pd.DataFrame, selected_period_ids: list[str], limit: int = 10) -> pd.DataFrame:
    ordered_ids = _ordered_period_ids(periods, selected_period_ids)
    if len(ordered_ids) < 2 or messages is None or messages.empty or "period_id" not in messages.columns:
        return pd.DataFrame()
    prev_id, cur_id = str(ordered_ids[-2]), str(ordered_ids[-1])
    prev_stats = build_tag_statistics(messages[messages["period_id"].astype(str) == prev_id].copy())
    cur_stats = build_tag_statistics(messages[messages["period_id"].astype(str) == cur_id].copy())
    if prev_stats.empty and cur_stats.empty:
        return pd.DataFrame()
    prev = prev_stats[["Тег", "Сообщений", "Охват", "Вовлеченность", "Негатив"]].copy() if not prev_stats.empty else pd.DataFrame(columns=["Тег", "Сообщений", "Охват", "Вовлеченность", "Негатив"])
    cur = cur_stats[["Тег", "Сообщений", "Охват", "Вовлеченность", "Негатив"]].copy() if not cur_stats.empty else pd.DataFrame(columns=["Тег", "Сообщений", "Охват", "Вовлеченность", "Негатив"])
    prev = prev.rename(columns={c: f"{c}_prev" for c in prev.columns if c != "Тег"})
    cur = cur.rename(columns={c: f"{c}_cur" for c in cur.columns if c != "Тег"})
    merged = cur.merge(prev, on="Тег", how="outer").fillna(0)
    for metric in ["Сообщений", "Охват", "Вовлеченность", "Негатив"]:
        merged[f"Δ {metric.lower()}"] = pd.to_numeric(merged.get(f"{metric}_cur", 0), errors="coerce").fillna(0) - pd.to_numeric(merged.get(f"{metric}_prev", 0), errors="coerce").fillna(0)
    merged["abs_delta"] = merged[["Δ сообщений", "Δ охват", "Δ вовлеченность", "Δ негатив"]].abs().sum(axis=1)
    merged = merged.sort_values("abs_delta", ascending=False).head(limit)
    out = pd.DataFrame({
        "Тег": merged["Тег"].astype(str),
        "Сообщений сейчас": merged["Сообщений_cur"].astype(int),
        "Δ сообщений": merged["Δ сообщений"].astype(int),
        "Охват сейчас": merged["Охват_cur"].astype(int),
        "Δ охвата": merged["Δ охват"].astype(int),
        "Вовлеченность сейчас": merged["Вовлеченность_cur"].astype(int),
        "Δ вовлеченности": merged["Δ вовлеченность"].astype(int),
        "Негатив сейчас": merged["Негатив_cur"].astype(int),
        "Δ негатива": merged["Δ негатив"].astype(int),
    })
    return out



def build_client_insights_summary(messages: pd.DataFrame, events_agg: pd.DataFrame, periods: pd.DataFrame, selected_period_ids: list[str], *, profile: str = "") -> str:
    """Return a text version of the client-insights block for automatic summaries."""
    metrics = overview_metrics(messages)
    sent = metrics.get("sentiment", {}) or {}
    total = int(sent.get("total", 0) or 0)
    negative = int(sent.get("negative", 0) or 0)
    negative_share = negative / total if total else 0.0
    engagement = int(metrics.get("engagement", 0) or 0)
    risk_level = "низкий" if negative_share < 0.01 else "средний" if negative_share < 0.05 else "высокий"

    lines: list[str] = []
    lines.append("Клиентский обзор")
    lines.append(f"Риск негатива: {risk_level}; негативных сообщений — {format_int(negative)} ({negative_share * 100:.1f}%).")
    lines.append(f"Суммарная вовлеченность: {format_int(engagement)}.")

    if len(selected_period_ids or []) >= 2:
        change_lines = build_period_change_insights(messages, periods, selected_period_ids)
        if change_lines:
            lines.append("Что изменилось к предыдущему периоду:")
            for item in change_lines[:5]:
                lines.append(f"• {item}")

        tag_changes = build_tag_change_table(messages, periods, selected_period_ids, limit=5)
        if tag_changes is not None and not tag_changes.empty:
            lines.append("Теги с заметными изменениями:")
            for _, row in tag_changes.head(5).iterrows():
                parts = [str(row.get("Тег") or "")]
                try:
                    delta_messages = int(row.get("Δ сообщений", 0) or 0)
                except Exception:
                    delta_messages = 0
                try:
                    delta_reach = int(row.get("Δ охвата", 0) or 0)
                except Exception:
                    delta_reach = 0
                try:
                    delta_eng = int(row.get("Δ вовлеченности", 0) or 0)
                except Exception:
                    delta_eng = 0
                details = []
                if delta_messages:
                    details.append(f"сообщения {format_int(delta_messages)}")
                if delta_reach:
                    details.append(f"охват {format_int(delta_reach)}")
                if delta_eng:
                    details.append(f"вовлеченность {format_int(delta_eng)}")
                if details:
                    parts.append("; ".join(details))
                lines.append("• " + " — ".join([p for p in parts if p]))

    tags = top_client_tags(messages, limit=5)
    if tags is not None and not tags.empty:
        lines.append("Топ тегов для отчета:")
        for _, row in tags.iterrows():
            lines.append(
                f"• {row.get('Тег', '')} — {format_int(row.get('Сообщений', 0))} сообщ.; "
                f"охват {format_int(row.get('Охват', 0))}; вовлеченность {format_int(row.get('Вовлеченность', 0))}."
            )

    top_events = top_client_events(events_agg, limit=5)
    if top_events is not None and not top_events.empty:
        title_col = _event_title_col(top_events) or "title"
        lines.append("Топ инфоповодов для отчета:")
        for _, row in top_events.iterrows():
            lines.append(f"• {row.get(title_col, '')} — {format_int(row.get('message_count', 0))} сообщ.")

    return "\n".join(line for line in lines if str(line).strip())

def render_client_insights(messages: pd.DataFrame, events_agg: pd.DataFrame, periods: pd.DataFrame, selected_period_ids: list[str], *, profile: str = "") -> None:
    st.subheader("Клиентский обзор")
    st.caption("Сводный слой для презентации заказчику: риски, ключевые сигналы и изменения между периодами.")

    metrics = overview_metrics(messages)
    sent = metrics.get("sentiment", {}) or {}
    total = int(sent.get("total", 0) or 0)
    negative = int(sent.get("negative", 0) or 0)
    negative_share = negative / total if total else 0.0
    engagement = int(metrics.get("engagement", 0) or 0)

    cards = st.columns(4)
    cards[0].metric("Риск негатива", "низкий" if negative_share < 0.01 else "средний" if negative_share < 0.05 else "высокий", help=f"Негативных сообщений: {format_int(negative)}")
    cards[1].metric("Доля негатива", f"{negative_share * 100:.1f}%")
    cards[2].metric("Вовлеченность", format_int(engagement))
    cards[3].metric("Инфоповодов", format_int(len(events_agg) if isinstance(events_agg, pd.DataFrame) else 0))

    signals: list[dict[str, Any]] = []
    if negative > 0:
        signals.append({"Сигнал": "Есть негативные сообщения", "Что смотреть": "Негативные публикации и темы с высокой вовлеченностью", "Данные": f"{format_int(negative)} сообщ. · {negative_share * 100:.1f}%", "Приоритет": "Средний" if negative_share < 0.05 else "Высокий"})
    else:
        signals.append({"Сигнал": "Критичный негатив не выявлен", "Что смотреть": "Контролировать всплески по тегам и инфоповодам", "Данные": "0 негативных сообщений", "Приоритет": "Низкий"})

    if isinstance(events_agg, pd.DataFrame) and not events_agg.empty and "negative_count" in events_agg.columns:
        risky_events = events_agg.copy()
        risky_events["negative_count"] = pd.to_numeric(risky_events["negative_count"], errors="coerce").fillna(0)
        risky_events = risky_events[(risky_events["negative_count"] > 0) & (~risky_events.get("title", pd.Series(dtype=str)).apply(is_technical_event_title))]
        if not risky_events.empty:
            risky_events = risky_events.sort_values(["negative_count", "message_count"], ascending=False).head(3)
            title_col = _event_title_col(risky_events) or "title"
            names = "; ".join(risky_events[title_col].astype(str).head(3).tolist())
            signals.append({"Сигнал": "Темы с негативом", "Что смотреть": names, "Данные": f"{format_int(risky_events['negative_count'].sum())} нег. сообщ.", "Приоритет": "Средний"})

    tags = build_tag_statistics(messages)
    if tags is not None and not tags.empty and "Негатив" in tags.columns:
        neg_tags = tags.copy()
        neg_tags["Негатив"] = pd.to_numeric(neg_tags["Негатив"], errors="coerce").fillna(0)
        neg_tags = neg_tags[neg_tags["Негатив"] > 0].sort_values(["Негатив", "Сообщений"], ascending=False).head(3)
        if not neg_tags.empty:
            tag_names = "; ".join(neg_tags["Тег"].astype(str).tolist())
            signals.append({"Сигнал": "Теги с негативом", "Что смотреть": tag_names, "Данные": f"{format_int(neg_tags['Негатив'].sum())} нег. сообщ.", "Приоритет": "Средний"})

    st.markdown("#### Риски и сигналы")
    st.dataframe(pd.DataFrame(signals), hide_index=True, use_container_width=True)

    if len(selected_period_ids or []) >= 2:
        st.markdown("#### Что изменилось к предыдущему периоду")
        insights = build_period_change_insights(messages, periods, selected_period_ids)
        if insights:
            for item in insights[:6]:
                st.markdown(f"- {item}")
        else:
            st.caption("Значимых изменений по основным метрикам не найдено.")
        tag_changes = build_tag_change_table(messages, periods, selected_period_ids, limit=10)
        if not tag_changes.empty:
            display = tag_changes.copy()
            for col in [c for c in display.columns if c != "Тег"]:
                display[col] = display[col].apply(format_int)
            with st.expander("Теги с наибольшими изменениями", expanded=True):
                st.dataframe(display, hide_index=True, use_container_width=True)

    st.markdown("#### Что включить в отчет")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Топ тегов**")
        top_tags = top_client_tags(messages, limit=5)
        if top_tags.empty:
            st.caption("Теги не найдены.")
        else:
            for _, row in top_tags.iterrows():
                st.markdown(f"- **{row['Тег']}** — {format_int(row.get('Сообщений', 0))} сообщ. · охват {format_int(row.get('Охват', 0))}")
    with c2:
        st.markdown("**Топ инфоповодов**")
        top_events = top_client_events(events_agg, limit=5)
        if top_events.empty:
            st.caption("Инфоповоды не найдены.")
        else:
            title_col = _event_title_col(top_events) or "title"
            for _, row in top_events.iterrows():
                st.markdown(f"- **{row.get(title_col, '')}** — {format_int(row.get('message_count', 0))} сообщ.")


def render_project_access(is_admin: bool) -> tuple[str | None, str, pd.DataFrame]:
    projects = list_projects(include_inactive=is_admin)
    if projects.empty:
        if is_admin:
            st.info("Пока нет проектов. Создайте первый проект в блоке управления ниже.")
        else:
            st.warning("Пока нет доступных проектов.")
        return None, "owner" if is_admin else "none", projects

    if is_admin:
        options = projects["project_id"].astype(str).tolist()
        labels = {str(r["project_id"]): str(r.get("project_name") or r["project_id"]) for _, r in projects.iterrows()}
        selected = st.sidebar.selectbox("Проект", options, format_func=lambda x: labels.get(x, x))
        return selected, "owner", projects

    if st.session_state.get("platform_project_id"):
        project_id = st.session_state["platform_project_id"]
        role = st.session_state.get("platform_project_role", "viewer")
        project_row = projects[projects["project_id"].astype(str) == str(project_id)]
        project_name = str(project_row.iloc[0].get("project_name") if not project_row.empty else project_id)
        st.sidebar.success(f"Доступ: {project_name} · {role}")
        if st.sidebar.button("Сменить проект / выйти"):
            st.session_state.pop("platform_project_id", None)
            st.session_state.pop("platform_project_role", None)
            st.rerun()
        return project_id, role, projects

    st.sidebar.info("Введите код доступа к проекту")
    access_code = st.sidebar.text_input("Код проекта", type="password", key="project_access_code")
    if st.sidebar.button("Открыть проект"):
        project_id, role = resolve_project_access(access_code)
        if project_id:
            st.session_state["platform_project_id"] = project_id
            st.session_state["platform_project_role"] = role
            st.rerun()
        else:
            st.sidebar.error("Код проекта не найден.")
    return None, "none", projects


def render_project_manager(projects: pd.DataFrame) -> None:
    st.header("Управление проектами")
    with st.expander("Создать проект", expanded=projects.empty):
        name = st.text_input("Название проекта", key="new_project_name")
        description = st.text_area("Описание проекта", key="new_project_description")
        topic_profile = st.selectbox(
            "Профиль алгоритма",
            list(ALGORITHM_PROFILE_OPTIONS.keys()),
            format_func=lambda x: ALGORITHM_PROFILE_OPTIONS.get(x, x),
            key="new_topic_profile",
            help="Профиль не привязывает платформу к одной отрасли: по умолчанию темы берутся из колонок выгрузки и универсальных правил.",
        )
        viewer_code = st.text_input("Код просмотра", type="password", key="new_viewer_code")
        editor_code = st.text_input("Код редактора", type="password", key="new_editor_code")
        if st.button("Создать проект", type="primary"):
            if not name.strip():
                st.error("Укажите название проекта.")
            else:
                project_id = create_project(
                    project_name=name,
                    description=description,
                    viewer_code=viewer_code,
                    editor_code=editor_code,
                    settings={"topic_profile": topic_profile},
                )
                st.success(f"Проект создан: {project_id}")
                st.rerun()

    if projects.empty:
        return
    st.subheader("Существующие проекты")
    view = projects.copy()
    for col in ["created_at", "updated_at"]:
        if col in view.columns:
            view[col] = view[col].apply(fmt_date)
    show = view[[c for c in ["project_name", "description", "status", "created_at", "updated_at", "project_id"] if c in view.columns]].rename(columns={
        "project_name": "Проект",
        "description": "Описание",
        "status": "Статус",
        "created_at": "Создан",
        "updated_at": "Обновлен",
        "project_id": "ID",
    })
    event = st.dataframe(show, hide_index=True, use_container_width=True, selection_mode="single-row", on_select="rerun")
    rows = getattr(event, "selection", {}).get("rows", []) if event is not None else []
    if rows:
        row = projects.iloc[rows[0]]
        project_id = str(row["project_id"])
        with st.expander(f"Редактировать: {row.get('project_name')}", expanded=True):
            new_name = st.text_input("Название", value=str(row.get("project_name") or ""), key=f"edit_project_name_{project_id}")
            new_description = st.text_area("Описание", value=str(row.get("description") or ""), key=f"edit_project_description_{project_id}")
            new_status = st.selectbox("Статус", ["active", "hidden", "archived"], index=["active", "hidden", "archived"].index(str(row.get("status") or "active")) if str(row.get("status") or "active") in ["active", "hidden", "archived"] else 0, key=f"edit_project_status_{project_id}")
            current_settings = project_settings_from_row(row)
            current_profile = str(current_settings.get("topic_profile") or "universal")
            if current_profile not in ALGORITHM_PROFILE_OPTIONS:
                current_profile = "universal"
            new_topic_profile = st.selectbox(
                "Профиль алгоритма",
                list(ALGORITHM_PROFILE_OPTIONS.keys()),
                index=list(ALGORITHM_PROFILE_OPTIONS.keys()).index(current_profile),
                format_func=lambda x: ALGORITHM_PROFILE_OPTIONS.get(x, x),
                key=f"edit_topic_profile_{project_id}",
            )
            current_chart_labels = chart_label_settings_from_project_settings(current_settings)
            with st.expander("Настройки подписей на графиках", expanded=False):
                chart_font = st.selectbox(
                    "Шрифт значений",
                    CHART_LABEL_FONT_OPTIONS,
                    index=CHART_LABEL_FONT_OPTIONS.index(current_chart_labels.get("font", "Arial")) if current_chart_labels.get("font", "Arial") in CHART_LABEL_FONT_OPTIONS else 0,
                    key=f"chart_label_font_{project_id}",
                )
                chart_font_size = st.slider(
                    "Размер шрифта",
                    min_value=8,
                    max_value=28,
                    value=int(current_chart_labels.get("font_size", 11)),
                    step=1,
                    key=f"chart_label_size_{project_id}",
                )
                position_keys = list(CHART_LABEL_POSITION_OPTIONS.keys())
                current_position = current_chart_labels.get("position", "top")
                chart_position = st.radio(
                    "Местоположение значений",
                    position_keys,
                    index=position_keys.index(current_position) if current_position in position_keys else 0,
                    format_func=lambda x: CHART_LABEL_POSITION_OPTIONS.get(x, x),
                    horizontal=True,
                    key=f"chart_label_position_{project_id}",
                )
                show_donut_legend = st.checkbox(
                    "Показывать легенду круговой диаграммы",
                    value=bool(current_chart_labels.get("show_donut_legend", False)),
                    help="По умолчанию легенда скрыта, потому что рядом с круговой диаграммой уже есть блок значений.",
                    key=f"chart_label_show_donut_legend_{project_id}",
                )
                st.caption("Настройка применяется к подписям на линейных, столбчатых и круговых графиках сравнения периодов.")

            current_branding = report_branding_from_project_settings(current_settings, project_name=str(row.get("project_name") or ""))
            with st.expander("Брендирование отчетов", expanded=False):
                report_client_name = st.text_input(
                    "Название клиента в отчете",
                    value=str(current_branding.get("client_name") or ""),
                    key=f"report_client_name_{project_id}",
                    help="Это название будет отображаться на титульной инфографике и в Word/PDF.",
                )
                report_title = st.text_input(
                    "Заголовок отчета",
                    value=str(current_branding.get("report_title") or "Дайджест упоминаний"),
                    key=f"report_title_{project_id}",
                )
                b1, b2 = st.columns(2)
                with b1:
                    report_accent_color = st.color_picker(
                        "Акцентный цвет",
                        value=_valid_hex_color(current_branding.get("accent_color"), "#2563eb"),
                        key=f"report_accent_color_{project_id}",
                    )
                with b2:
                    report_background_color = st.color_picker(
                        "Фон инфографики",
                        value=_valid_hex_color(current_branding.get("background_color"), "#ffffff"),
                        key=f"report_background_color_{project_id}",
                    )
                report_footer_text = st.text_input(
                    "Подпись в футере",
                    value=str(current_branding.get("footer_text") or ""),
                    key=f"report_footer_text_{project_id}",
                    help="Например: подготовлено агентством / внутренний аналитический отчет.",
                )
                report_logo_url = st.text_input(
                    "URL логотипа (резерв под следующий этап)",
                    value=str(current_branding.get("logo_url") or ""),
                    key=f"report_logo_url_{project_id}",
                    help="Поле сохраняется в настройках проекта. Визуальная вставка логотипа будет подключена отдельным этапом.",
                )
                st.caption("Брендирование применяется к Word/PDF/PNG-выгрузкам саммари и клиентских отчетов.")

            st.caption("Коды доступа заполняйте только если хотите заменить текущие.")
            new_viewer_code = st.text_input("Новый код просмотра", type="password", key=f"edit_viewer_code_{project_id}")
            new_editor_code = st.text_input("Новый код редактора", type="password", key=f"edit_editor_code_{project_id}")
            if st.button("Сохранить проект", key=f"save_project_{project_id}"):
                updated_settings = dict(current_settings)
                updated_settings["topic_profile"] = new_topic_profile
                updated_settings["chart_label_settings"] = {
                    "font": chart_font,
                    "font_size": int(chart_font_size),
                    "position": chart_position,
                    "show_donut_legend": bool(show_donut_legend),
                }
                updated_settings["report_branding"] = {
                    "client_name": report_client_name,
                    "report_title": report_title,
                    "accent_color": report_accent_color,
                    "background_color": report_background_color,
                    "footer_text": report_footer_text,
                    "logo_url": report_logo_url,
                }
                update_project(
                    project_id,
                    project_name=new_name,
                    description=new_description,
                    status=new_status,
                    viewer_code=new_viewer_code,
                    editor_code=new_editor_code,
                    settings=updated_settings,
                )
                st.success("Проект обновлен.")
                st.rerun()


def render_period_selector(project_id: str) -> tuple[list[str], pd.DataFrame]:
    periods = list_periods(project_id, include_inactive=False)
    if periods.empty:
        st.sidebar.warning("В проекте пока нет загруженных периодов.")
        return [], periods
    labels = {}
    for _, r in periods.iterrows():
        period_id = str(r["period_id"])
        labels[period_id] = f"{r.get('period_name') or period_id} · {fmt_period(r)}"
    default = periods["period_id"].astype(str).head(3).tolist()
    selected = st.sidebar.multiselect("Периоды", periods["period_id"].astype(str).tolist(), default=default, format_func=lambda x: labels.get(x, x))
    return selected, periods


def read_uploaded_to_canonical(uploaded_file, source_system: str) -> pd.DataFrame:
    suffix = Path(uploaded_file.name).suffix.lower() or ".csv"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        tmp_path = tmp.name
    try:
        return read_source_table(tmp_path, source_system=source_system)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def read_generated_tables_from_dir(output_dir: Path) -> dict[str, pd.DataFrame]:
    return {name: read_table(str(output_dir), name) for name in ["events", "discussions", "messages", "discussion_messages", "event_discussions"]}


def render_upload_page(project_id: str, role: str, work_dir: str) -> None:
    st.header("Загрузка файла")
    if role_rank(role) < role_rank("editor"):
        st.info("Для загрузки файлов нужен доступ редактора или владельца.")
        return

    with st.form("upload_form"):
        period_name = st.text_input("Название периода", placeholder="Например: 24.04.2026–30.04.2026")
        date_col1, date_col2 = st.columns(2)
        with date_col1:
            date_from = st.date_input("Дата начала", value=None, format="DD.MM.YYYY")
        with date_col2:
            date_to = st.date_input("Дата окончания", value=None, format="DD.MM.YYYY")
        source_system = st.selectbox("Источник", ["auto", "mediologia", "mediologia_excel", "brand_analytics", "generic"], format_func=lambda x: {
            "auto": "Автоопределение",
            "mediologia": "Медиалогия CSV",
            "mediologia_excel": "Медиалогия Excel",
            "brand_analytics": "Brand Analytics",
            "generic": "Универсальный CSV/Excel",
        }.get(x, x))
        uploaded = st.file_uploader("CSV или Excel", type=["csv", "xlsx", "xls", "xlsm"])
        st.caption("Алгоритм")
        c1, c2, c3 = st.columns(3)
        with c1:
            threshold = st.slider("Похожесть", 0.10, 0.60, 0.30, 0.01)
        with c2:
            event_gap_hours = st.slider("Разрыв между волнами, часов", 1.0, 24.0, 3.0, 1.0)
        with c3:
            event_window_hours = st.slider("Макс. окно инфоповода, часов", 4.0, 72.0, 16.0, 4.0)
        submitted = st.form_submit_button("Обработать и сохранить", type="primary")

    if not submitted:
        return
    if uploaded is None:
        st.error("Загрузите файл.")
        return
    if not period_name.strip():
        st.error("Укажите название периода.")
        return
    if date_from and date_to and date_from > date_to:
        st.error("Дата начала не может быть позже даты окончания.")
        return

    with st.spinner("Читаю файл и привожу к единому формату..."):
        try:
            canonical = read_uploaded_to_canonical(uploaded, source_system)
        except Exception as exc:
            st.error("Не удалось прочитать файл.")
            st.info(
                "Проверьте, что файл содержит лист/таблицу с сообщениями: дата, текст/сообщение, url/ссылка, источник или автор. "
                "Если в Excel несколько листов, платформа автоматически ищет лист «Сообщения» и пропускает пустые листы."
            )
            st.exception(exc)
            return
    st.success(f"Файл прочитан: {len(canonical):,} строк".replace(",", " "))
    with st.expander("Предпросмотр распознанных колонок", expanded=False):
        st.dataframe(canonical.head(20), use_container_width=True)

    period_id = make_period_id(project_id, period_name, uploaded.name)
    output_dir = Path(work_dir) / project_id / period_id
    output_dir.mkdir(parents=True, exist_ok=True)
    with st.spinner("Собираю сообщения, обсуждения и инфоповоды..."):
        manifest = run_preprocess_from_dataframe(
            canonical,
            output=output_dir,
            source_file=uploaded.name,
            similarity_threshold=float(threshold),
            event_gap_hours=float(event_gap_hours),
            event_window_hours=float(event_window_hours),
        )
        tables = read_generated_tables_from_dir(output_dir)

    storage_path = ""
    try:
        storage_path = save_uploaded_file_to_storage(project_id, period_id, uploaded.name, uploaded.getvalue())
    except Exception as exc:
        st.warning(f"Обработанные данные сохраню в БД, но сырой файл не удалось сохранить в Storage: {exc}")

    with st.spinner("Сохраняю данные проекта в Supabase..."):
        manifest = dict(manifest or {})
        manifest.update({"storage_path": storage_path, "source_system": source_system})
        save_processed_tables(
            project_id=project_id,
            period_id=period_id,
            period_name=period_name,
            source_filename=uploaded.name,
            tables=tables,
            manifest=manifest,
            date_from=date_from,
            date_to=date_to,
            replace=True,
        )
    st.success("Период сохранен в платформенной базе.")
    clear_platform_caches(project_id)


def render_period_history(project_id: str, role: str) -> None:
    st.header("История периодов")
    if role_rank(role) < role_rank("editor"):
        st.info("Для редактирования истории нужен доступ редактора или владельца.")
        return
    periods = list_periods(project_id, include_inactive=True)
    if periods.empty:
        st.info("Периодов пока нет.")
        return
    view = periods.copy()
    view["Период"] = view.apply(fmt_period, axis=1)
    show = view[[c for c in ["period_name", "Период", "source_filename", "status", "period_id"] if c in view.columns]].rename(columns={
        "period_name": "Название",
        "source_filename": "Файл",
        "status": "Статус",
        "period_id": "ID",
    })
    event = st.dataframe(show, hide_index=True, use_container_width=True, selection_mode="single-row", on_select="rerun")
    rows = getattr(event, "selection", {}).get("rows", []) if event is not None else []
    if not rows:
        return
    row = periods.iloc[rows[0]]
    period_id = str(row["period_id"])
    with st.expander("Редактировать период", expanded=True):
        name = st.text_input("Название периода", value=str(row.get("period_name") or ""), key=f"period_name_{period_id}")
        def to_date(v):
            ts = pd.to_datetime(v, errors="coerce")
            return None if pd.isna(ts) else ts.date()
        date_from = st.date_input("Дата начала", value=to_date(row.get("date_from")), format="DD.MM.YYYY", key=f"date_from_{period_id}")
        date_to = st.date_input("Дата окончания", value=to_date(row.get("date_to")), format="DD.MM.YYYY", key=f"date_to_{period_id}")
        filename = st.text_input("Файл", value=str(row.get("source_filename") or ""), key=f"filename_{period_id}")
        status = st.selectbox("Статус", ["active", "hidden", "archived"], index=["active", "hidden", "archived"].index(str(row.get("status") or "active")) if str(row.get("status") or "active") in ["active", "hidden", "archived"] else 0, key=f"status_{period_id}")
        comment = st.text_area("Комментарий", value=str((row.get("manifest") or {}).get("comment", "")) if isinstance(row.get("manifest"), dict) else "", key=f"comment_{period_id}")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("Сохранить изменения", key=f"save_period_{period_id}"):
                if date_from and date_to and date_from > date_to:
                    st.error("Дата начала не может быть позже даты окончания.")
                else:
                    update_period_metadata(project_id, period_id, period_name=name, date_from=date_from, date_to=date_to, source_filename=filename, status=status, manifest_updates={"comment": comment})
                    st.success("Период обновлен.")
                    st.rerun()
        with c2:
            if st.button("Скрыть период", key=f"hide_period_{period_id}"):
                delete_period(project_id, period_id, hard=False)
                st.success("Период скрыт.")
                st.rerun()

        with st.expander("Удалить выгрузку без восстановления", expanded=False):
            st.warning(
                "Удаление выгрузки удалит период и все обработанные таблицы этого периода из platform_table_rows. "
                "Также будут удалены ручные правки, которые явно ссылаются на этот период. "
                "Другие проекты и другие периоды не затрагиваются."
            )
            manifest = row.get("manifest") if isinstance(row.get("manifest"), dict) else {}
            storage_path = str((manifest or {}).get("storage_path") or "").strip()
            if storage_path:
                st.caption(f"Исходный файл в Storage: {storage_path}")
            delete_storage = st.checkbox(
                "Удалить исходный файл из Supabase Storage, если он был сохранен",
                value=True,
                key=f"delete_storage_{period_id}",
            )
            st.caption("Удаление запускается одной кнопкой. Действие необратимо.")
            if st.button("Удалить выгрузку", key=f"hard_delete_period_{period_id}", type="primary"):
                try:
                    result = delete_period(
                        project_id,
                        period_id,
                        hard=True,
                        delete_storage=delete_storage,
                        cleanup_manual=True,
                    )
                except Exception as exc:
                    st.error("Не удалось удалить выгрузку.")
                    st.exception(exc)
                    return

                manual_count = int(result.get("manual_rows_deleted") or 0) if isinstance(result, dict) else 0
                table_count = int(result.get("table_rows_deleted") or 0) if isinstance(result, dict) else 0
                storage_deleted = bool(result.get("storage_deleted")) if isinstance(result, dict) else False
                mode = str(result.get("mode") or "") if isinstance(result, dict) else ""
                for warning in (result.get("warnings") or []) if isinstance(result, dict) else []:
                    st.warning(str(warning))
                if delete_storage and storage_path and not storage_deleted and mode != "soft_fallback":
                    st.warning("Выгрузка удалена из базы, но исходный файл в Storage удалить не удалось или он уже отсутствовал.")
                if mode == "soft_fallback":
                    st.warning("Физическое удаление не завершилось, поэтому период скрыт из интерфейса. Для полной очистки можно повторить удаление позже или выполнить очистку в Supabase.")
                else:
                    st.success(f"Выгрузка удалена. Удалено строк данных: {table_count}. Удалено связанных ручных правок: {manual_count}.")
                clear_platform_caches(project_id)
                st.rerun()

@st.cache_data(show_spinner=False)
def enrich_messages(messages: pd.DataFrame, event_discussions: pd.DataFrame, discussion_messages: pd.DataFrame, events: pd.DataFrame) -> pd.DataFrame:
    """Attach event ids/titles to messages safely.

    Some project profiles, especially Brand Analytics story-based imports, may
    have generated event rows but no discussion/event link table in older
    periods. The previous version assumed that the merge through
    discussion_messages/event_discussions always created an `event_id` column and
    crashed with KeyError when it did not.
    """
    if messages is None or messages.empty:
        return messages if isinstance(messages, pd.DataFrame) else pd.DataFrame()

    out = messages.copy()
    if "message_id" not in out.columns:
        out["message_id"] = out.index.astype(str)
    out["message_id"] = out["message_id"].fillna("").astype(str)

    # 1) Preferred path: message -> discussion -> event links.
    if (
        isinstance(event_discussions, pd.DataFrame)
        and isinstance(discussion_messages, pd.DataFrame)
        and not event_discussions.empty
        and not discussion_messages.empty
        and "discussion_id" in event_discussions.columns
        and "discussion_id" in discussion_messages.columns
    ):
        link = discussion_messages.merge(event_discussions, on="discussion_id", how="left")
        if "message_id" in link.columns and "event_id" in link.columns:
            msg_event = (
                link[["message_id", "event_id"]]
                .dropna(subset=["message_id"])
                .drop_duplicates("message_id")
            )
            msg_event["message_id"] = msg_event["message_id"].fillna("").astype(str)
            out = out.merge(msg_event, on="message_id", how="left", suffixes=("", "_linked"))
            if "event_id_linked" in out.columns:
                if "event_id" not in out.columns:
                    out["event_id"] = out["event_id_linked"]
                else:
                    out["event_id"] = out["event_id"].fillna(out["event_id_linked"])
                out = out.drop(columns=["event_id_linked"])

    # 2) Fallback for story-based BA periods: map source topic/story to event.
    if "event_id" not in out.columns:
        out["event_id"] = ""
    out["event_id"] = out["event_id"].fillna("").astype(str)

    if isinstance(events, pd.DataFrame) and not events.empty and "event_id" in events.columns:
        events_work = events.copy()
        events_work["event_id"] = events_work["event_id"].fillna("").astype(str)

        needs_fallback = out["event_id"].str.strip().eq("").all()
        if needs_fallback:
            topic_map: dict[str, str] = {}
            if "source_main_topic" in events_work.columns:
                for _, row in events_work.dropna(subset=["event_id"]).iterrows():
                    key = str(row.get("source_main_topic") or "").strip().lower()
                    if key and key not in topic_map:
                        topic_map[key] = str(row.get("event_id"))
            if "event_title" in events_work.columns:
                for _, row in events_work.dropna(subset=["event_id"]).iterrows():
                    key = str(row.get("event_title") or "").strip().lower()
                    if key and key not in topic_map:
                        topic_map[key] = str(row.get("event_id"))

            if topic_map:
                source_col = None
                for candidate in ["source_main_topic", "source_topics", "event_title"]:
                    if candidate in out.columns:
                        source_col = candidate
                        break
                if source_col:
                    keys = out[source_col].fillna("").astype(str).str.split(";").str[0].str.strip().str.lower()
                    out["event_id"] = keys.map(topic_map).fillna(out["event_id"])

        # Add/fill event title without requiring a merge key to exist.
        if "event_title" in events_work.columns:
            title_map = (
                events_work.drop_duplicates("event_id")
                .set_index("event_id")["event_title"]
                .fillna("")
                .astype(str)
                .to_dict()
            )
            mapped_titles = out["event_id"].fillna("").astype(str).map(title_map).fillna("")
            if "event_title" not in out.columns:
                out["event_title"] = mapped_titles
            else:
                current = out["event_title"].fillna("").astype(str)
                out["event_title"] = current.where(current.str.strip().ne(""), mapped_titles)

    if "event_title" not in out.columns:
        out["event_title"] = ""
    return out


@st.cache_data(show_spinner=False)
def aggregate_events(events: pd.DataFrame) -> pd.DataFrame:
    if events.empty:
        return events
    df = events.copy()
    df["title"] = df.get("event_title", "").fillna("Без названия").astype(str).replace("", "Без названия")
    df["group_key"] = df["title"].str.lower().str.strip()
    rows = []
    for key, group in df.groupby("group_key", dropna=False):
        row = {
            "group_key": key,
            "title": group["title"].iloc[0],
            "description": pick_event_description(group),
            "tags": " | ".join(sorted(set("|".join(group.get("main_tags", pd.Series(dtype=str)).fillna("").astype(str)).split("|")) - {""})),
            "start_date": pd.to_datetime(group.get("start_date"), errors="coerce").min(),
            "end_date": pd.to_datetime(group.get("end_date"), errors="coerce").max(),
            "message_count": int(pd.to_numeric(group.get("message_count", 0), errors="coerce").fillna(0).sum()),
            "chat_count": int(pd.to_numeric(group.get("chat_count", 0), errors="coerce").fillna(0).sum()),
            "negative_count": int(pd.to_numeric(group.get("negative_count", 0), errors="coerce").fillna(0).sum()),
            "importance_score": float(pd.to_numeric(group.get("importance_score", 0), errors="coerce").fillna(0).max()),
            "event_ids": list(group["event_id"].astype(str)) if "event_id" in group.columns else [],
        }
        row["negative_share"] = row["negative_count"] / row["message_count"] if row["message_count"] else 0
        rows.append(row)
    out = pd.DataFrame(rows)
    if not out.empty:
        out = out.sort_values(["importance_score", "message_count"], ascending=False)
    return out


def build_event_description(group: pd.DataFrame) -> str:
    """Build a neutral, source-agnostic event description."""
    text = " ".join(group.get("event_summary", pd.Series(dtype=str)).fillna("").astype(str).tolist())
    tags = " | ".join(sorted(set("|".join(group.get("main_tags", pd.Series(dtype=str)).fillna("").astype(str)).split("|")) - {""}))
    low = f"{text} {tags}".lower().replace("ё", "е")
    patterns = [
        ("проблемы, жалобы и негативный опыт", ["жалоб", "проблем", "негатив", "ошиб", "не работает", "плохо", "брак", "дефект"]),
        ("цены, стоимость и условия", ["цен", "стоим", "скид", "акци", "тариф", "услов", "дорого", "дешев"]),
        ("качество продукта или услуги", ["качеств", "материал", "характерист", "свойств", "надежн", "эффектив"]),
        ("наличие, поставки и логистика", ["достав", "налич", "склад", "постав", "логист", "срок", "отгруз"]),
        ("монтаж, применение и эксплуатация", ["монтаж", "установ", "примен", "использ", "эксплуатац", "строител", "утепл", "изоляц"]),
        ("документы, сертификаты и требования", ["сертифик", "документ", "декларац", "гост", "снип", "требован", "стандарт"]),
        ("безопасность и риски", ["безопас", "пожар", "огне", "горюч", "опасн", "токсич"]),
        ("экология и энергоэффективность", ["эколог", "энергоэфф", "энергосбереж", "устойчив", "переработ"]),
        ("конкуренты и сравнение", ["конкур", "аналог", "сравнен", "рынок", "бренд"]),
        ("клиентский сервис и поддержка", ["поддерж", "сервис", "менеджер", "дилер", "магазин", "клиент"]),
    ]
    signals = []
    for label, keys in patterns:
        if any(k in low for k in keys):
            signals.append(label)
    if signals:
        return "В теме обсуждались: " + "; ".join(signals[:5]) + "."
    return f"В теме обсуждались: {tags}." if tags else "В теме обсуждались связанные сообщения выбранного периода."


def pick_event_description(group: pd.DataFrame) -> str:
    """Return manual description if present; otherwise build an automatic one."""
    for col in ["display_description", "manual_description", "event_description"]:
        if col in group.columns:
            vals = [str(x).strip() for x in group[col].fillna("").tolist() if str(x).strip()]
            if vals:
                return vals[0]
    return build_event_description(group)


def manual_payloads(manual_df: pd.DataFrame, table_name: str) -> list[dict[str, Any]]:
    if manual_df is None or manual_df.empty or "table_name" not in manual_df.columns:
        return []
    rows = manual_df[manual_df["table_name"].astype(str) == table_name]
    payloads: list[dict[str, Any]] = []
    for _, row in rows.iterrows():
        payload = row.get("payload") or {}
        if isinstance(payload, dict):
            payload = dict(payload)
            payload.setdefault("_row_key", str(row.get("row_key") or ""))
            payloads.append(payload)
    return payloads


def get_manual_state(project_id: str) -> dict[str, Any]:
    try:
        manual_df = list_manual(project_id)
    except Exception:
        manual_df = pd.DataFrame()

    hidden_messages: set[str] = set()
    hidden_message_keys: dict[str, str] = {}
    for payload in manual_payloads(manual_df, "message_hidden"):
        message_id = str(payload.get("message_id") or payload.get("_row_key", "").replace("message_hidden::", ""))
        if message_id:
            hidden_messages.add(message_id)
            hidden_message_keys[message_id] = str(payload.get("_row_key") or f"message_hidden::{message_id}")

    irrelevant_pairs: set[tuple[str, str]] = set()
    irrelevant_keys: dict[tuple[str, str], str] = {}
    for payload in manual_payloads(manual_df, "message_irrelevant"):
        event_id = str(payload.get("event_id") or "")
        message_id = str(payload.get("message_id") or "")
        if event_id and message_id:
            pair = (event_id, message_id)
            irrelevant_pairs.add(pair)
            irrelevant_keys[pair] = str(payload.get("_row_key") or f"message_irrelevant::{event_id}::{message_id}")

    move_map: dict[str, str] = {}
    for payload in manual_payloads(manual_df, "message_moves"):
        message_id = str(payload.get("message_id") or payload.get("_row_key", "").replace("message_move::", ""))
        target_event_id = str(payload.get("target_event_id") or "")
        if message_id and target_event_id:
            move_map[message_id] = target_event_id

    event_edits: dict[str, dict[str, Any]] = {}
    for payload in manual_payloads(manual_df, "event_edits"):
        event_id = str(payload.get("event_id") or payload.get("_row_key", "").replace("event_edit::", ""))
        if event_id:
            event_edits[event_id] = payload

    event_merges: dict[str, str] = {}
    for payload in manual_payloads(manual_df, "event_merges"):
        source_event_id = str(payload.get("source_event_id") or payload.get("_row_key", "").replace("event_merge::", ""))
        target_event_id = str(payload.get("target_event_id") or "")
        if source_event_id and target_event_id:
            event_merges[source_event_id] = target_event_id

    manual_events = manual_payloads(manual_df, "manual_events")

    return {
        "manual_df": manual_df,
        "hidden_messages": hidden_messages,
        "hidden_message_keys": hidden_message_keys,
        "irrelevant_pairs": irrelevant_pairs,
        "irrelevant_keys": irrelevant_keys,
        "move_map": move_map,
        "event_edits": event_edits,
        "event_merges": event_merges,
        "manual_events": manual_events,
    }


def append_manual_events(events: pd.DataFrame, manual_events: list[dict[str, Any]]) -> pd.DataFrame:
    if not manual_events:
        return events
    rows = []
    for payload in manual_events:
        event_id = str(payload.get("event_id") or "").strip()
        title = str(payload.get("title") or "").strip()
        if not event_id or not title:
            continue
        rows.append({
            "event_id": event_id,
            "event_title": title,
            "event_summary": str(payload.get("description") or ""),
            "display_description": str(payload.get("description") or ""),
            "main_tags": str(payload.get("tags") or "Ручной инфоповод"),
            "start_date": pd.NaT,
            "end_date": pd.NaT,
            "message_count": 0,
            "chat_count": 0,
            "negative_count": 0,
            "importance_score": 0,
            "status": str(payload.get("status") or "active"),
            "is_manual_event": True,
        })
    if not rows:
        return events
    extra = pd.DataFrame(rows)
    if events is None or events.empty:
        return extra
    return pd.concat([events, extra], ignore_index=True, sort=False)


def recompute_event_counts(events: pd.DataFrame, messages: pd.DataFrame) -> pd.DataFrame:
    if events is None or events.empty or messages is None or messages.empty or "event_id" not in messages.columns:
        return events
    out = events.copy()
    msg = messages.copy()
    msg["event_id"] = msg["event_id"].fillna("").astype(str)
    msg = msg[msg["event_id"].str.strip() != ""]
    if msg.empty:
        return out
    grouped = msg.groupby("event_id", dropna=False)
    for event_id, group in grouped:
        mask = out["event_id"].astype(str) == str(event_id)
        if not mask.any():
            continue
        out.loc[mask, "message_count"] = int(len(group))
        if "chat_title" in group.columns:
            out.loc[mask, "chat_count"] = int(group["chat_title"].fillna("").astype(str).replace("", pd.NA).dropna().nunique())
        elif "chat_id" in group.columns:
            out.loc[mask, "chat_count"] = int(group["chat_id"].fillna("").astype(str).replace("", pd.NA).dropna().nunique())
        if "sentiment" in group.columns:
            out.loc[mask, "negative_count"] = int(group["sentiment"].fillna("").astype(str).str.lower().str.contains("нег").sum())
        if "datetime" in group.columns:
            dt = pd.to_datetime(group["datetime"], errors="coerce").dropna()
            if not dt.empty:
                out.loc[mask, "start_date"] = dt.min()
                out.loc[mask, "end_date"] = dt.max()
    return out


@st.cache_data(show_spinner=False)
def apply_manual_overrides(project_id: str, events: pd.DataFrame, messages: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    state = get_manual_state(project_id)
    events_out = events.copy() if events is not None else pd.DataFrame()
    messages_out = messages.copy() if messages is not None else pd.DataFrame()

    events_out = append_manual_events(events_out, state["manual_events"])
    if not events_out.empty:
        for col in ["event_id", "event_title", "event_summary", "main_tags", "status", "display_description"]:
            if col not in events_out.columns:
                events_out[col] = ""

    # Apply event-level manual edits.
    if not events_out.empty:
        for event_id, payload in state["event_edits"].items():
            mask = events_out["event_id"].astype(str) == str(event_id)
            if not mask.any():
                continue
            if str(payload.get("title") or "").strip():
                events_out.loc[mask, "event_title"] = str(payload.get("title")).strip()
            if "description" in payload:
                events_out.loc[mask, "display_description"] = str(payload.get("description") or "").strip()
                events_out.loc[mask, "event_summary"] = str(payload.get("description") or "").strip()
            if str(payload.get("tags") or "").strip():
                events_out.loc[mask, "main_tags"] = str(payload.get("tags")).strip()
            if str(payload.get("status") or "").strip():
                events_out.loc[mask, "status"] = str(payload.get("status")).strip()

    # Merge events by redirecting source title to target title.
    if not events_out.empty and state["event_merges"]:
        title_map = {
            str(row.get("event_id")): str(row.get("event_title") or "Без названия")
            for _, row in events_out.iterrows()
        }
        summary_map = {
            str(row.get("event_id")): str(row.get("display_description") or row.get("event_summary") or "")
            for _, row in events_out.iterrows()
        }
        for source_event_id, target_event_id in state["event_merges"].items():
            mask = events_out["event_id"].astype(str) == str(source_event_id)
            if mask.any():
                events_out.loc[mask, "event_title"] = title_map.get(str(target_event_id), str(target_event_id))
                events_out.loc[mask, "display_description"] = summary_map.get(str(target_event_id), "")
                events_out.loc[mask, "merged_into"] = str(target_event_id)

    if not messages_out.empty:
        if "message_id" not in messages_out.columns:
            messages_out["message_id"] = messages_out.index.astype(str)
        messages_out["message_id"] = messages_out["message_id"].fillna("").astype(str)

        if state["hidden_messages"]:
            messages_out = messages_out[~messages_out["message_id"].isin(state["hidden_messages"])].copy()

        if "event_id" in messages_out.columns:
            messages_out["event_id"] = messages_out["event_id"].fillna("").astype(str)
            if state["event_merges"]:
                messages_out["event_id"] = messages_out["event_id"].map(lambda x: state["event_merges"].get(str(x), str(x)))
            if state["move_map"]:
                messages_out["event_id"] = messages_out.apply(lambda r: state["move_map"].get(str(r.get("message_id")), str(r.get("event_id") or "")), axis=1)

    # Refresh titles in messages after moves/merges.
    if not events_out.empty and not messages_out.empty and "event_id" in messages_out.columns:
        title_map = {
            str(row.get("event_id")): str(row.get("event_title") or "Без названия")
            for _, row in events_out.iterrows()
        }
        messages_out["event_title"] = messages_out["event_id"].fillna("").astype(str).map(title_map).fillna(messages_out.get("event_title", ""))

    events_out = recompute_event_counts(events_out, messages_out)

    # Hide events after counts are recomputed.
    if not events_out.empty and "status" in events_out.columns:
        events_out = events_out[~events_out["status"].fillna("").astype(str).str.lower().isin(["hidden", "deleted", "archived"])].copy()

    return events_out, messages_out, state


def create_manual_event(project_id: str, title: str, description: str = "", tags: str = "", status: str = "active") -> str:
    event_id = "manual_" + uuid.uuid4().hex[:12]
    save_manual(project_id, "manual_events", f"manual_event::{event_id}", {
        "event_id": event_id,
        "title": title.strip(),
        "description": description.strip(),
        "tags": tags.strip() or "Ручной инфоповод",
        "status": status,
    })
    return event_id


def event_select_options(events_agg: pd.DataFrame, exclude_event_ids: set[str] | None = None) -> list[tuple[str, str]]:
    exclude_event_ids = exclude_event_ids or set()
    options: list[tuple[str, str]] = []
    if events_agg is None or events_agg.empty:
        return options
    for _, row in events_agg.iterrows():
        ids = [str(x) for x in row.get("event_ids", [])]
        if not ids:
            continue
        if set(ids) & exclude_event_ids:
            continue
        label = f"{row.get('title') or ids[0]} · {int(row.get('message_count') or 0)} сообщ."
        options.append((ids[0], label))
    return options


def build_auto_summary(messages: pd.DataFrame, events_agg: pd.DataFrame, periods: pd.DataFrame, selected_period_ids: list[str]) -> str:
    total = len(messages)
    chats = messages["chat_title"].nunique() if "chat_title" in messages.columns else 0
    authors = messages["author"].nunique() if "author" in messages.columns else 0
    neg = 0
    if "sentiment" in messages.columns:
        neg = int(messages["sentiment"].fillna("").astype(str).str.lower().str.contains("нег").sum())
    neg_share = (neg / total * 100) if total else 0
    period_names = []
    if not periods.empty:
        subset = periods[periods["period_id"].astype(str).isin([str(x) for x in selected_period_ids])]
        period_names = [str(x) for x in subset.get("period_name", pd.Series(dtype=str)).tolist()]
    top_events = events_agg.head(5)["title"].tolist() if not events_agg.empty else []
    top_chats = []
    if "chat_title" in messages.columns:
        top_chats = messages["chat_title"].fillna("").astype(str).replace("", pd.NA).dropna().value_counts().head(5).index.tolist()

    lines = []
    lines.append(f"За выбранный период обработано {total:,} сообщений из {chats:,} чатов; уникальных авторов — {authors:,}.".replace(",", " "))
    lines.append(f"Негативных сообщений: {neg:,} ({neg_share:.1f}%).".replace(",", " "))
    if period_names:
        lines.append("Периоды: " + "; ".join(period_names[:6]) + ("…" if len(period_names) > 6 else ""))
    if top_events:
        lines.append("Основные инфоповоды: " + "; ".join(top_events) + ".")
    if top_chats:
        lines.append("Наиболее активные чаты: " + "; ".join(top_chats) + ".")

    client_overview = build_client_insights_summary(messages, events_agg, periods, selected_period_ids)
    if client_overview:
        lines.append(client_overview)
    return "\n\n".join(lines)


def summary_storage_key(period_ids: list[str], profile: str = "") -> str:
    prefix = "summary"
    if is_taxi_project_profile(profile):
        prefix = "summary::taxi"
    return prefix + "::" + "__".join(sorted(str(x) for x in (period_ids or []) if str(x).strip()))


def render_period_summary(
    project_id: str,
    project_name: str,
    period_ids: list[str],
    messages: pd.DataFrame,
    events_agg: pd.DataFrame,
    periods: pd.DataFrame,
    role: str,
    *,
    profile: str = "",
    metrics: dict[str, Any] | None = None,
    branding: dict[str, Any] | None = None,
) -> None:
    """Unified editable/exportable period summary for all project profiles."""
    st.subheader("Саммари периода")
    key = summary_storage_key(period_ids, profile)
    manual = get_manual(project_id, key)
    auto_summary = build_taxi_auto_summary(messages, events_agg, periods, period_ids) if is_taxi_project_profile(profile) else build_auto_summary(messages, events_agg, periods, period_ids)
    summary_text = str((manual or {}).get("summary") or "").strip() or auto_summary
    st.markdown(summary_text.replace("\n", "  \n"))

    metrics = metrics or overview_metrics(messages)
    metrics.setdefault("period_label", selected_period_label(periods, period_ids))
    metrics.setdefault("project_name", project_name)
    period_label = str(metrics.get("period_label") or selected_period_label(periods, period_ids))

    with st.expander("Выгрузить саммари", expanded=False):
        st.caption("Можно скачать Word, PDF или отдельную PNG-инфографику. В инфографику попадут метрики, тональность, топ-теги, топ-инфоповоды и ключевые тезисы саммари.")
        render_summary_export_buttons(project_name, period_label, summary_text, metrics, key_prefix=f"summary_export_{abs(hash(key))}", messages=messages, events_agg=events_agg, branding=branding)

    if role_rank(role) >= role_rank("editor"):
        with st.expander("Редактировать саммари", expanded=False):
            edited = st.text_area("Текст саммари", value=summary_text, height=220, key=f"summary_{key}")
            c1, c2 = st.columns(2)
            with c1:
                if st.button("Сохранить саммари", key=f"save_{key}"):
                    save_manual(project_id, "summaries", key, {"summary": edited, "period_ids": period_ids, "profile": profile})
                    st.success("Саммари сохранено.")
                    st.rerun()
            with c2:
                if st.button("Вернуть автоматическое", key=f"auto_{key}"):
                    delete_manual(project_id, key)
                    st.success("Вернули автоматическое саммари.")
                    st.rerun()


def render_summary(project_id: str, period_ids: list[str], messages: pd.DataFrame, events_agg: pd.DataFrame, periods: pd.DataFrame, role: str) -> None:
    """Backward-compatible wrapper for older calls."""
    render_period_summary(project_id, "Проект", period_ids, messages, events_agg, periods, role)


def render_period_dynamics(messages: pd.DataFrame, periods: pd.DataFrame, period_ids: list[str]) -> None:
    if len(period_ids) < 2 or messages.empty or "period_id" not in messages.columns:
        return
    rows = []
    for period_id, group in messages.groupby("period_id"):
        meta = periods[periods["period_id"].astype(str) == str(period_id)]
        name = str(meta.iloc[0].get("period_name") if not meta.empty else period_id)
        sort_date = pd.to_datetime(meta.iloc[0].get("date_from") if not meta.empty else None, errors="coerce")
        neg = int(group.get("sentiment", pd.Series(dtype=str)).fillna("").astype(str).str.lower().str.contains("нег").sum()) if "sentiment" in group.columns else 0
        total = len(group)
        rows.append({"period_name": name, "sort_date": sort_date, "Сообщения": total, "Негатив": neg, "Доля негатива, %": round(neg / total * 100, 1) if total else 0})
    summary = pd.DataFrame(rows)
    if summary.empty:
        return
    summary = summary.sort_values(["sort_date", "period_name"], na_position="last")
    st.subheader("Динамика по периодам")
    chart_type = st.selectbox(
        "Тип визуализации динамики",
        ["График", "Столбчатая", "Круговая диаграмма"],
        index=0,
        key=f"period_dynamics_chart_type_{abs(hash(tuple(summary['period_name'].astype(str).tolist())))}",
    )
    chart_df = summary.copy()
    if chart_type == "Круговая диаграмма":
        metric = st.selectbox(
            "Метрика для круговой диаграммы динамики",
            ["Сообщения", "Негатив"],
            index=0,
            key=f"period_dynamics_pie_metric_{abs(hash(tuple(summary['period_name'].astype(str).tolist())))}",
        )
        _render_value_distribution_donut(chart_df.rename(columns={"period_name": "Период"}), "Период", metric, metric)
    elif chart_type == "Столбчатая":
        c1, c2 = st.columns(2)
        with c1:
            bars = alt.Chart(chart_df).mark_bar(size=70).encode(
                x=alt.X("period_name:N", sort=None, title="Период", axis=alt.Axis(labelAngle=-90)),
                y=alt.Y("Сообщения:Q", title="Сообщения"),
                tooltip=["period_name", alt.Tooltip("Сообщения:Q", format=",")],
            )
            st.altair_chart(bars.properties(height=300), use_container_width=True)
        with c2:
            neg_bars = alt.Chart(chart_df).mark_bar(size=70).encode(
                x=alt.X("period_name:N", sort=None, title="Период", axis=alt.Axis(labelAngle=-90)),
                y=alt.Y("Доля негатива, %:Q", title="Доля негатива, %"),
                tooltip=["period_name", alt.Tooltip("Доля негатива, %:Q", format=".1f")],
            )
            st.altair_chart(neg_bars.properties(height=300), use_container_width=True)
    else:
        c1, c2 = st.columns(2)
        with c1:
            line = alt.Chart(chart_df).mark_line(point=True).encode(
                x=alt.X("period_name:N", sort=None, title="Период"),
                y=alt.Y("Сообщения:Q", title="Сообщения"),
                tooltip=["period_name", alt.Tooltip("Сообщения:Q", format=",")],
            )
            st.altair_chart(line.properties(height=300), use_container_width=True)
        with c2:
            neg_line = alt.Chart(chart_df).mark_line(point=True).encode(
                x=alt.X("period_name:N", sort=None, title="Период"),
                y=alt.Y("Доля негатива, %:Q", title="Доля негатива, %"),
                tooltip=["period_name", alt.Tooltip("Доля негатива, %:Q", format=".1f")],
            )
            st.altair_chart(neg_line.properties(height=300), use_container_width=True)
    st.dataframe(summary.drop(columns=["sort_date"]), hide_index=True, use_container_width=True)



def selected_event_filter_key(project_id: str | None) -> str:
    return f"selected_event_filter::{project_id or 'global'}"


def set_selected_event_filter(project_id: str | None, selected: pd.Series) -> None:
    event_ids = [str(x) for x in (selected.get("event_ids", []) or []) if str(x).strip()]
    if not event_ids and "event_id" in selected.index:
        event_ids = [str(selected.get("event_id"))]
    st.session_state[selected_event_filter_key(project_id)] = {
        "title": str(selected.get("title") or selected.get("event_title") or "Выбранный инфоповод"),
        "event_ids": event_ids,
        "group_key": str(selected.get("group_key") or ""),
    }


def get_selected_event_filter(project_id: str | None) -> dict[str, Any] | None:
    value = st.session_state.get(selected_event_filter_key(project_id))
    return value if isinstance(value, dict) and value.get("event_ids") else None


def clear_selected_event_filter(project_id: str | None) -> None:
    st.session_state.pop(selected_event_filter_key(project_id), None)


def filter_messages_by_selected_event(messages: pd.DataFrame, event_filter: dict[str, Any] | None) -> pd.DataFrame:
    if messages is None or messages.empty or not event_filter:
        return messages
    event_ids = {str(x) for x in event_filter.get("event_ids", []) if str(x).strip()}
    if not event_ids:
        return messages
    mask = pd.Series(False, index=messages.index)
    for col in ["event_id", "source_event_id", "final_event_id", "source_final_event_id"]:
        if col in messages.columns:
            mask = mask | messages[col].fillna("").astype(str).isin(event_ids)
    if mask.any():
        return messages[mask].copy()

    # Fallback for imported sources where event links may be reconstructed by title.
    title = str(event_filter.get("title") or "").strip().lower()
    if title:
        for col in ["event_title", "source_main_topic", "Сюжет"]:
            if col in messages.columns:
                fallback_mask = messages[col].fillna("").astype(str).str.strip().str.lower().eq(title)
                if fallback_mask.any():
                    return messages[fallback_mask].copy()
    return messages.iloc[0:0].copy()


def _event_series_filter(selected: pd.Series) -> dict[str, Any]:
    event_ids = [str(x) for x in (selected.get("event_ids", []) or []) if str(x).strip()]
    if not event_ids and "event_id" in selected.index:
        event_ids = [str(selected.get("event_id"))]
    return {
        "title": str(selected.get("title") or selected.get("event_title") or "Выбранный инфоповод"),
        "event_ids": event_ids,
        "group_key": str(selected.get("group_key") or selected.get("event_id") or selected.get("title") or "event"),
    }


def _event_tags_text(selected: pd.Series, event_messages: pd.DataFrame) -> str:
    values: list[str] = []
    for col in ["tags", "main_tags", "display_tags", "source_topics"]:
        if col in selected.index:
            values.extend(split_pipe_values(str(selected.get(col) or "")))
    if not values and isinstance(event_messages, pd.DataFrame) and not event_messages.empty and "tags" in event_messages.columns:
        for item in event_messages["tags"].fillna("").astype(str).head(500).tolist():
            values.extend(split_pipe_values(item))
    seen: set[str] = set()
    clean: list[str] = []
    for value in values:
        key = value.strip().lower().replace("ё", "е")
        if value.strip() and key not in seen:
            seen.add(key)
            clean.append(value.strip())
    return ", ".join(clean[:12])


def _event_auto_summary(selected: pd.Series, event_messages: pd.DataFrame) -> str:
    description = str(selected.get("description") or selected.get("display_description") or selected.get("summary") or "").strip()
    if description:
        return description
    title = str(selected.get("title") or selected.get("event_title") or "инфоповод").strip()
    count = int(len(event_messages)) if isinstance(event_messages, pd.DataFrame) else int(selected.get("message_count", 0) or 0)
    negative_count = 0
    if isinstance(event_messages, pd.DataFrame) and not event_messages.empty:
        sent = sentiment_counts(event_messages)
        negative_count = int(sent.get("negative", 0) or 0)
    return f"В теме «{title}» собрано {format_int(count)} сообщений. Негативных сообщений: {format_int(negative_count)}."


def render_selected_event_detail(project_id: str, selected: pd.Series, messages: pd.DataFrame) -> None:
    """Render a unified selected-infopoint card across all project profiles."""
    event_filter = _event_series_filter(selected)
    event_messages = filter_messages_by_selected_event(messages, event_filter)
    title = str(selected.get("title") or selected.get("event_title") or "Выбранный инфоповод")
    summary = _event_auto_summary(selected, event_messages)
    tags_text = _event_tags_text(selected, event_messages)

    st.markdown(f"## {title}")
    if summary:
        st.info(summary)

    metrics = overview_metrics(event_messages if isinstance(event_messages, pd.DataFrame) else pd.DataFrame())
    sent = metrics.get("sentiment", {}) or {}
    total = int(sent.get("total", metrics.get("messages", 0)) or 0)
    chat_count = int(selected.get("chat_count", 0) or 0)
    if not chat_count and isinstance(event_messages, pd.DataFrame) and not event_messages.empty:
        for col in ["chat_title", "platform", "source", "Источник", "Место публикации"]:
            if col in event_messages.columns:
                chat_count = int(event_messages[col].fillna("").astype(str).replace("", pd.NA).dropna().nunique())
                break
    author_count = 0
    if isinstance(event_messages, pd.DataFrame) and not event_messages.empty:
        for col in ["author", "Автор"]:
            if col in event_messages.columns:
                author_count = int(event_messages[col].fillna("").astype(str).replace("", pd.NA).dropna().nunique())
                break

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Сообщений", format_int(metrics.get("messages", 0)))
    c2.metric("Источников/чатов", format_int(chat_count))
    c3.metric("Авторов", format_int(author_count))
    c4.metric("Негатив", percent_text(int(sent.get("negative", 0) or 0), total))
    c5.metric("Важность", str(round(float(selected.get("importance_score", 0) or 0), 2)))

    m1, m2, m3 = st.columns(3)
    m1.metric("Аудитория", format_int(metrics.get("audience", 0)))
    m2.metric("Охват", format_int(metrics.get("reach", 0)))
    m3.metric("Вовлеченность", format_int(metrics.get("engagement", 0)))

    if tags_text:
        st.caption(f"Теги: {tags_text}")
    if isinstance(event_messages, pd.DataFrame):
        st.caption(f"В выбранном инфоповоде найдено сообщений: {format_int(len(event_messages))}.")

    mode = st.radio(
        "Сообщения инфоповода",
        ["Ключевые сообщения", "Вся лента"],
        horizontal=True,
        key=f"selected_event_messages_mode_{project_id}_{abs(hash(str(event_filter.get('group_key'))))}",
    )

    if event_messages is None or event_messages.empty:
        st.info("Сообщений по выбранному инфоповоду не найдено.")
        return

    work = event_messages.copy()
    text_col = message_text_column(work)
    link_col = message_link_column(work)
    work["_audience"] = numeric_series(work, ["audience", "Аудитория"]).astype(int)
    work["_reach"] = numeric_series(work, ["views", "Просмотры", "Просмотров", "reach", "Охват"]).astype(int)
    work["_engagement"] = numeric_series(work, ["engagement", "Вовлечённость", "Вовлеченность", "engagement_count"]).astype(int)

    if mode == "Ключевые сообщения":
        st.caption("Показаны топ-15 сообщений выбранного инфоповода по вовлеченности. Если вовлеченность равна 0, учитываются охват и аудитория.")
        view = work.sort_values(["_engagement", "_reach", "_audience"], ascending=False).head(15).copy()
    else:
        search_key = f"selected_event_feed_search_{project_id}_{abs(hash(str(event_filter.get('group_key'))))}"
        search = st.text_input("Поиск по ленте инфоповода", placeholder="Введите слово или фразу", key=search_key)
        view = work.copy()
        if search.strip() and text_col:
            view = view[view[text_col].fillna("").astype(str).str.contains(search.strip(), case=False, regex=False)]
        view = view.sort_values("datetime", ascending=False) if "datetime" in view.columns else view
        total_found = int(len(view))
        page_size = int(st.selectbox(
            "Сообщений на странице",
            [25, 50, 100, 200],
            index=1,
            key=f"selected_event_feed_page_size_{project_id}_{abs(hash(str(event_filter.get('group_key'))))}",
        ))
        total_pages = max(1, (total_found + page_size - 1) // page_size)
        page = int(st.number_input(
            "Страница",
            min_value=1,
            max_value=total_pages,
            value=1,
            step=1,
            key=f"selected_event_feed_page_{project_id}_{abs(hash(str(event_filter.get('group_key'))))}",
        ))
        start = (page - 1) * page_size
        end = start + page_size
        st.caption(
            f"Найдено сообщений: {format_int(total_found)}. "
            f"Показано: {format_int(start + 1 if total_found else 0)}–{format_int(min(end, total_found))} из {format_int(total_found)}."
        )
        view = view.iloc[start:end].copy()

    _render_message_list(view, text_col=text_col, link_col=link_col)


def render_events(
    project_id: str,
    role: str,
    events_agg: pd.DataFrame,
    messages: pd.DataFrame,
    manual_state: dict[str, Any],
) -> None:
    st.subheader("Инфоповоды")
    can_edit = role_rank(role) >= role_rank("editor")

    if can_edit:
        with st.expander("Создать инфоповод вручную", expanded=False):
            title = st.text_input("Название нового инфоповода", key="new_manual_event_title")
            description = st.text_area("Описание", key="new_manual_event_description")
            tags = st.text_input("Теги", key="new_manual_event_tags")
            if st.button("Создать инфоповод", type="primary", key="create_manual_event"):
                if not title.strip():
                    st.error("Укажите название инфоповода.")
                else:
                    create_manual_event(project_id, title, description, tags)
                    st.success("Инфоповод создан.")
                    st.rerun()

    if events_agg.empty:
        st.info("Инфоповоды не найдены.")
        return

    word = st.text_input("Фильтр по слову в сообщениях", placeholder="Например: доставка, качество, сертификат")
    filtered_events = events_agg.copy()
    filtered_messages = messages.copy()

    text_col = message_text_column(filtered_messages)
    if word.strip() and text_col:
        mask = filtered_messages[text_col].fillna("").astype(str).str.contains(word.strip(), case=False, regex=False)
        filtered_messages = filtered_messages[mask]
        if "event_id" in filtered_messages.columns:
            allowed = set(filtered_messages["event_id"].dropna().astype(str))
            filtered_events = filtered_events[filtered_events["event_ids"].apply(lambda ids: bool(set(map(str, ids)) & allowed))]
        st.caption(f"Найдено сообщений: {len(filtered_messages):,}".replace(",", " "))
        msg_view = filtered_messages.copy()
        if not msg_view.empty:
            msg_view["Дата"] = msg_view.get("datetime", "").apply(fmt_date)
            if text_col:
                msg_view["Текст"] = msg_view[text_col].fillna("").astype(str).str.slice(0, 700)
            link_col = message_link_column(msg_view)
            msg_view["Ссылка"] = msg_view[link_col].fillna("").astype(str) if link_col else ""
            columns = [c for c in ["Дата", "chat_title", "author", "event_title", "Текст", "Ссылка"] if c in msg_view.columns]
            st.dataframe(
                msg_view[columns].rename(columns={
                    "chat_title": "Источник/площадка",
                    "author": "Автор",
                    "event_title": "Инфоповод",
                }).head(500),
                hide_index=True,
                use_container_width=True,
                column_config={"Ссылка": st.column_config.LinkColumn("Ссылка")},
            )

    table = filtered_events.copy()
    table["Период"] = table.apply(
        lambda r: f"{fmt_date(r.get('start_date'))}–{fmt_date(r.get('end_date'))}"
        if fmt_date(r.get("start_date")) != fmt_date(r.get("end_date"))
        else fmt_date(r.get("start_date")),
        axis=1,
    )
    table["Негатив"] = (table["negative_share"] * 100).round(1).astype(str) + "%"
    show = table[["title", "description", "Период", "message_count", "chat_count", "Негатив", "importance_score"]].rename(columns={
        "title": "Сюжет / инфоповод",
        "description": "Описание",
        "message_count": "Сообщений",
        "chat_count": "Источников",
        "importance_score": "Важность",
    })
    event = st.dataframe(show, hide_index=True, use_container_width=True, selection_mode="single-row", on_select="rerun")

    rows = getattr(event, "selection", {}).get("rows", []) if event is not None else []
    if not rows:
        return

    selected = filtered_events.iloc[rows[0]]
    set_selected_event_filter(project_id, selected)
    selected_ids = set(map(str, selected.get("event_ids", [])))
    render_selected_event_detail(project_id, selected, messages)

    if can_edit:
        with st.expander("Правка выбранного инфоповода", expanded=False):
            new_title = st.text_input("Название", value=str(selected.get("title") or ""), key=f"edit_title_{selected.get('group_key')}")
            new_desc = st.text_area("Описание", value=str(selected.get("description") or ""), height=160, key=f"edit_desc_{selected.get('group_key')}")
            new_tags = st.text_input("Теги", value=str(selected.get("tags") or ""), key=f"edit_tags_{selected.get('group_key')}")
            c1, c2, c3 = st.columns(3)
            with c1:
                if st.button("Сохранить правки", key=f"save_event_edit_{selected.get('group_key')}"):
                    for event_id in selected_ids:
                        save_manual(project_id, "event_edits", f"event_edit::{event_id}", {
                            "event_id": event_id,
                            "title": new_title,
                            "description": new_desc,
                            "tags": new_tags,
                            "status": "active",
                        })
                    st.success("Правки сохранены.")
                    st.rerun()
            with c2:
                if st.button("Скрыть инфоповод", key=f"hide_event_{selected.get('group_key')}"):
                    for event_id in selected_ids:
                        save_manual(project_id, "event_edits", f"event_edit::{event_id}", {
                            "event_id": event_id,
                            "title": new_title,
                            "description": new_desc,
                            "tags": new_tags,
                            "status": "hidden",
                        })
                    st.success("Инфоповод скрыт.")
                    st.rerun()
            with c3:
                options = event_select_options(events_agg, exclude_event_ids=selected_ids)
                if options:
                    target = st.selectbox("Объединить с темой", options, format_func=lambda x: x[1], key=f"merge_target_{selected.get('group_key')}")
                    if st.button("Объединить", key=f"merge_event_{selected.get('group_key')}"):
                        target_event_id = target[0]
                        for source_event_id in selected_ids:
                            if source_event_id != target_event_id:
                                save_manual(project_id, "event_merges", f"event_merge::{source_event_id}", {
                                    "source_event_id": source_event_id,
                                    "target_event_id": target_event_id,
                                })
                        st.success("Инфоповоды объединены.")
                        st.rerun()

    st.caption("Этот же инфоповод сохранен как фильтр для общего раздела «Ключевые сообщения / Вся лента».")



def _value_from_row(row: pd.Series, *columns: str) -> str:
    """Return the first non-empty value from a message row."""
    for col in columns:
        if col in row.index:
            value = str(row.get(col) or "").strip()
            if value and value.lower() not in {"nan", "none", "nat", "null"}:
                return value
    return ""


def _render_message_list(view: pd.DataFrame, *, text_col: str | None, link_col: str | None) -> None:
    """Render messages as readable cards instead of a dataframe."""
    if view is None or view.empty:
        st.info("Сообщений для показа нет.")
        return

    for _, row in view.iterrows():
        date_text = fmt_date(row.get("datetime")) if "datetime" in row.index else ""
        source = _value_from_row(row, "chat_title", "platform", "source", "Источник", "Место публикации")
        author = _value_from_row(row, "author", "Автор")
        sentiment = _value_from_row(row, "sentiment", "Тональность")
        event_title = _value_from_row(row, "event_title", "source_main_topic", "Сюжет")
        tags = _value_from_row(row, "tags", "Теги").replace("|", ", ")
        audience = int(row.get("_audience", 0) or 0)
        reach = int(row.get("_reach", 0) or 0)
        engagement = int(row.get("_engagement", 0) or 0)
        text = str(row.get(text_col, "") or "").strip() if text_col else ""
        link = str(row.get(link_col, "") or "").strip() if link_col else ""

        meta_parts = [part for part in [date_text, source, author, sentiment] if part]
        metrics_parts = [f"аудитория: {format_int(audience)}", f"охват: {format_int(reach)}", f"вовлеченность: {format_int(engagement)}"]

        st.markdown("---")
        if meta_parts:
            st.caption(" · ".join(meta_parts))
        if event_title:
            st.markdown(f"**Инфоповод:** {event_title}")
        if tags:
            st.caption(f"Теги: {tags}")
        st.markdown(f"*{' · '.join(metrics_parts)}*")
        st.write(text[:1800] if text else "—")
        if link.startswith("http"):
            st.markdown(f"[Открыть сообщение]({link})")


def render_messages_block(messages: pd.DataFrame, *, project_id: str | None = None) -> None:
    """Render key messages and full feed as a readable list.

    If an event was selected in the «Инфоповоды» section, both modes are
    filtered by that event: top messages and the full feed show only messages
    from the selected infopoint.
    """
    st.subheader("Ключевые сообщения")
    if messages is None or messages.empty:
        st.info("Сообщения не найдены.")
        return

    event_filter = get_selected_event_filter(project_id)
    if event_filter:
        c1, c2 = st.columns([4, 1])
        with c1:
            st.info(f"Выбран инфоповод: {event_filter.get('title')}. В топе и общей ленте показаны только сообщения этого инфоповода.")
        with c2:
            if st.button("Сбросить фильтр", key=f"clear_event_message_filter_{project_id or 'global'}", use_container_width=True):
                clear_selected_event_filter(project_id)
                st.rerun()

    mode = st.radio(
        "Режим просмотра сообщений",
        ["Ключевые сообщения", "Вся лента"],
        horizontal=True,
        key="messages_block_mode",
    )

    work = messages.copy()
    if event_filter:
        work = filter_messages_by_selected_event(work, event_filter)
        if work.empty:
            st.warning("По выбранному инфоповоду сообщения не найдены. Возможно, данные были пересобраны или связи инфоповодов изменились.")
            return
    text_col = message_text_column(work)
    link_col = message_link_column(work)
    work["_audience"] = numeric_series(work, ["audience", "Аудитория"]).astype(int)
    work["_reach"] = numeric_series(work, ["views", "Просмотры", "Просмотров", "reach", "Охват"]).astype(int)
    work["_engagement"] = numeric_series(work, ["engagement", "Вовлечённость", "Вовлеченность", "engagement_count"]).astype(int)

    if mode == "Ключевые сообщения":
        scope = "выбранного инфоповода" if event_filter else "всей выборки"
        st.caption(f"Показаны 15 сообщений с максимальной вовлеченностью для {scope}. Если вовлеченность равна 0, дополнительными критериями выступают охват и аудитория.")
        view = work.sort_values(["_engagement", "_reach", "_audience"], ascending=False).head(15).copy()
    else:
        search = st.text_input("Поиск по всей ленте", placeholder="Введите слово или фразу", key="full_feed_search")
        view = work.copy()
        if search.strip() and text_col:
            view = view[view[text_col].fillna("").astype(str).str.contains(search.strip(), case=False, regex=False)]
        view = view.sort_values("datetime", ascending=False) if "datetime" in view.columns else view

        total_found = int(len(view))
        page_size = int(st.selectbox("Сообщений на странице", [25, 50, 100, 200], index=1, key="full_feed_page_size"))
        total_pages = max(1, (total_found + page_size - 1) // page_size)
        page = int(st.number_input("Страница", min_value=1, max_value=total_pages, value=min(int(st.session_state.get("full_feed_page", 1)), total_pages), step=1, key="full_feed_page"))
        start = (page - 1) * page_size
        end = start + page_size
        st.caption(
            f"Найдено сообщений: {format_int(total_found)}. "
            f"Показано: {format_int(start + 1 if total_found else 0)}–{format_int(min(end, total_found))} из {format_int(total_found)}."
        )
        view = view.iloc[start:end].copy()

    _render_message_list(view, text_col=text_col, link_col=link_col)



# -----------------------------------------------------------------------------
# Driver chats dashboard profile
# -----------------------------------------------------------------------------


def taxi_bool_negative(messages: pd.DataFrame) -> pd.Series:
    if messages is None or messages.empty:
        return pd.Series(dtype=bool)
    if "is_negative" in messages.columns:
        return messages["is_negative"].astype(str).str.lower().isin(["true", "1", "yes", "да", "негатив", "negative"])
    if "sentiment" in messages.columns:
        return messages["sentiment"].fillna("").astype(str).str.lower().str.contains("нег|negative|отриц", regex=True, na=False)
    return pd.Series([False] * len(messages), index=messages.index)


def taxi_bool_positive(messages: pd.DataFrame) -> pd.Series:
    if messages is None or messages.empty:
        return pd.Series(dtype=bool)
    if "sentiment" in messages.columns:
        return messages["sentiment"].fillna("").astype(str).str.lower().str.contains("позит|positive|полож", regex=True, na=False)
    return pd.Series([False] * len(messages), index=messages.index)


def render_taxi_overview_statistics(events_agg: pd.DataFrame, messages: pd.DataFrame) -> None:
    """Top-level metrics for the driver-chat profile."""
    st.subheader("Статистика")
    total_messages = int(len(messages)) if isinstance(messages, pd.DataFrame) else 0
    chat_col = "chat_title" if isinstance(messages, pd.DataFrame) and "chat_title" in messages.columns else "chat_id" if isinstance(messages, pd.DataFrame) and "chat_id" in messages.columns else None
    author_col = "author" if isinstance(messages, pd.DataFrame) and "author" in messages.columns else "author_id" if isinstance(messages, pd.DataFrame) and "author_id" in messages.columns else None
    chat_count = int(messages[chat_col].fillna("").astype(str).replace("", pd.NA).dropna().nunique()) if chat_col and total_messages else 0
    author_count = int(messages[author_col].fillna("").astype(str).replace("", pd.NA).dropna().nunique()) if author_col and total_messages else 0
    neg_count = int(taxi_bool_negative(messages).sum()) if total_messages else 0
    high_count = 0
    if isinstance(events_agg, pd.DataFrame) and not events_agg.empty and "importance_score" in events_agg.columns:
        importance = pd.to_numeric(events_agg["importance_score"], errors="coerce").fillna(0)
        threshold = float(importance.quantile(0.75)) if len(importance) else 0
        high_count = int((importance >= threshold).sum()) if threshold else 0
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Инфоповодов", format_int(len(events_agg) if isinstance(events_agg, pd.DataFrame) else 0))
    c2.metric("Сообщений", format_int(total_messages))
    c3.metric("Чатов", format_int(chat_count))
    c4.metric("Негатив", f"{(neg_count / total_messages * 100):.0f}%" if total_messages else "0%")
    c5.metric("Высокая важность", format_int(high_count))
    if author_count:
        st.caption(f"Уникальных авторов: {format_int(author_count)}")


def normalize_taxi_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").lower().replace("ё", "е")).strip()


def taxi_macro_title(row: pd.Series) -> str:
    """Map detailed taxi events to report-level topics."""
    text = normalize_taxi_text(" ".join([
        str(row.get("event_title") or ""),
        str(row.get("event_summary") or ""),
        str(row.get("display_description") or ""),
        str(row.get("main_tags") or ""),
        str(row.get("microtopic") or ""),
    ]))
    rules = [
        ("Забастовка, бойкот и коллективные действия", ["strike", "забаст", "бойкот", "стачк", "митинг", "коллективн"]),
        ("Законы, налоги и регулирование такси", ["tax_law", "налог", "патент", "самозан", "минтранс", "реестр", "закон", "разрешен", "лиценз"]),
        ("Коэффициенты, приоритет и тарифы", ["coeff_priority", "коэфф", "коэф", "кэф", "приоритет", "тариф", "подач"]),
        ("Сбои и ошибки в приложении", ["app_bug", "сбой", "ошиб", "завис", "не работает", "яндекс про", "обновлен", "приложен"]),
        ("Проблемы с заказами в приложении", ["app_orders", "заказ", "назнач", "раздач", "цепоч", "не приход"]),
        ("Оплата, выплаты и удержания", ["payments", "оплат", "выплат", "деньг", "баланс", "удерж", "комисс"]),
        ("Блокировки и доступ к аккаунту", ["account_block", "блок", "аккаунт", "доступ", "вериф", "фотоконтроль"]),
        ("Детские кресла и требования к заказам", ["child_seat", "кресл", "детск", "ребен", "ребён"]),
        ("Карты, адреса и навигация", ["gps_map", "карт", "адрес", "навиг", "геолока", "gps", "маршрут"]),
        ("Заказы и правила в аэропортах", ["airport", "аэропорт", "шереметьево", "домодедово", "внуково", "пулково"]),
        ("Поддержка, парк и диспетчерские вопросы", ["support", "поддерж", "таксопарк", "диспетчер", "парк"]),
        ("Запуск и обсуждение WB Такси", ["wb_launch", "wb такси", "wildberries", "вайлдбер", "вб такси"]),
        ("Обсуждение сервиса Фастен", ["fasten", "фастен", "fasten_service"]),
        ("Общее обсуждение Яндекса", ["general_yandex", "яндекс", "yandex", "яша"]),
    ]
    for title, keys in rules:
        if any(key in text for key in keys):
            return title
    raw_title = str(row.get("event_title") or "").strip()
    return raw_title or "Прочие обсуждения"


def aggregate_taxi_events(events: pd.DataFrame, level: str = "balanced") -> pd.DataFrame:
    """Aggregate taxi events for three levels of detail."""
    if events is None or events.empty:
        return pd.DataFrame()
    df = events.copy()
    if "event_id" not in df.columns:
        df["event_id"] = df.index.astype(str)
    if "event_title" not in df.columns:
        df["event_title"] = "Без названия"
    if level == "detailed":
        df["__group_title"] = df["event_title"].fillna("Без названия").astype(str).replace("", "Без названия")
        df["__group_key"] = df["event_id"].astype(str)
    elif level == "macro":
        df["__group_title"] = df.apply(taxi_macro_title, axis=1)
        df["__group_key"] = df["__group_title"].map(normalize_taxi_text)
    else:
        df["__group_title"] = df["event_title"].fillna("Без названия").astype(str).replace("", "Без названия")
        df["__group_key"] = df["__group_title"].map(normalize_taxi_text)

    rows: list[dict[str, Any]] = []
    for key, group in df.groupby("__group_key", dropna=False):
        title = str(group["__group_title"].iloc[0] or "Без названия")
        tags = " | ".join(sorted(set("|".join(group.get("main_tags", pd.Series(dtype=str)).fillna("").astype(str)).split("|")) - {""}))
        msg_count = int(pd.to_numeric(group.get("message_count", 0), errors="coerce").fillna(0).sum())
        neg_count = int(pd.to_numeric(group.get("negative_count", 0), errors="coerce").fillna(0).sum())
        rows.append({
            "group_key": str(key),
            "title": title,
            "description": pick_event_description(group),
            "tags": tags,
            "start_date": pd.to_datetime(group.get("start_date"), errors="coerce").min(),
            "end_date": pd.to_datetime(group.get("end_date"), errors="coerce").max(),
            "message_count": msg_count,
            "chat_count": int(pd.to_numeric(group.get("chat_count", 0), errors="coerce").fillna(0).sum()),
            "negative_count": neg_count,
            "importance_score": float(pd.to_numeric(group.get("importance_score", 0), errors="coerce").fillna(0).max()),
            "event_ids": list(group["event_id"].astype(str)),
            "source_event_count": int(group["event_id"].astype(str).nunique()),
        })
    out = pd.DataFrame(rows)
    if out.empty:
        return out
    out["negative_share"] = out.apply(lambda r: float(r["negative_count"]) / float(r["message_count"]) if float(r.get("message_count") or 0) else 0.0, axis=1)
    return out.sort_values(["importance_score", "message_count"], ascending=False).reset_index(drop=True)


def build_taxi_auto_summary(messages: pd.DataFrame, events_agg: pd.DataFrame, periods: pd.DataFrame, selected_period_ids: list[str]) -> str:
    """Readable summary for driver-chat projects."""
    total = int(len(messages)) if isinstance(messages, pd.DataFrame) else 0
    if not total:
        return "По выбранному периоду пока нет сообщений для саммари."
    neg_count = int(taxi_bool_negative(messages).sum())
    pos_count = int(taxi_bool_positive(messages).sum())
    neutral_count = max(0, total - neg_count - pos_count)
    top_events = "нет выраженных тем"
    if isinstance(events_agg, pd.DataFrame) and not events_agg.empty:
        top = events_agg.sort_values("message_count", ascending=False).head(6)
        top_events = "; ".join(f"{r['title']} — {format_int(r['message_count'])}" for _, r in top.iterrows())
    chat_col = "chat_title" if "chat_title" in messages.columns else "chat_id" if "chat_id" in messages.columns else None
    top_chats = "нет данных"
    if chat_col:
        vc = messages[chat_col].fillna("").astype(str).replace("", pd.NA).dropna().value_counts().head(5)
        if not vc.empty:
            top_chats = "; ".join(f"{name} — {format_int(count)}" for name, count in vc.items())
    lines = [
        f"За выбранный период собрано {format_int(total)} сообщений. Тональность: {neutral_count / total * 100:.0f}% нейтрал, {neg_count / total * 100:.0f}% негатив, {pos_count / total * 100:.0f}% позитив.",
        f"Основные обсуждения: {top_events}.",
        f"Наиболее активные чаты: {top_chats}.",
    ]
    summary_text = "\n".join("• " + line for line in lines)
    client_overview = build_client_insights_summary(messages, events_agg, periods, selected_period_ids, profile="driver_chats")
    if client_overview:
        summary_text += "\n\n" + client_overview
    return summary_text


def render_taxi_summary(project_id: str, period_ids: list[str], messages: pd.DataFrame, events_agg: pd.DataFrame, periods: pd.DataFrame, role: str) -> None:
    key = "summary::taxi::" + "|".join(sorted(map(str, period_ids)))
    saved = get_manual(project_id, "summaries", key)
    auto = build_taxi_auto_summary(messages, events_agg, periods, period_ids)
    text = str((saved or {}).get("summary") or "").strip() or auto
    st.subheader("Саммари")
    st.markdown(text.replace("\n", "  \n"))
    if role_rank(role) >= role_rank("editor"):
        with st.expander("Редактировать саммари", expanded=False):
            edited = st.text_area("Саммари", value=text, height=220, key=f"taxi_summary_{key}")
            if st.button("Сохранить саммари", key=f"save_taxi_summary_{key}"):
                save_manual(project_id, "summaries", key, {"summary": edited, "period_ids": period_ids, "profile": "driver_chats"})
                st.success("Саммари сохранено.")
                st.rerun()


def render_taxi_dashboard(
    project_id: str,
    project_name: str,
    role: str,
    selected_period_ids: list[str],
    periods: pd.DataFrame,
    events: pd.DataFrame,
    messages: pd.DataFrame,
    manual_state: dict[str, Any],
    chart_label_settings: dict[str, Any] | None = None,
    report_branding: dict[str, Any] | None = None,
) -> None:
    """Dedicated UI for driver-chat digest projects inside the platform namespace."""
    level = st.sidebar.selectbox(
        "Уровень сборки инфоповодов",
        ["balanced", "macro", "detailed"],
        index=0,
        format_func=lambda x: {
            "balanced": "Сбалансировано — рекомендовано",
            "macro": "Крупные темы — для отчета",
            "detailed": "Подробно — первичные инфоповоды",
        }.get(x, x),
        key="taxi_event_detail_level",
    )
    raw_events_agg = aggregate_taxi_events(events, level=level)
    min_event_messages = render_min_event_messages_control("driver_chats", raw_events_agg, key="taxi_min_event_messages")
    events_agg, hidden_events, hidden_messages = filter_small_events(raw_events_agg, min_event_messages)
    metrics = render_project_intro(
        project_name,
        messages,
        periods,
        selected_period_ids,
        profile_label="Дайджест водительских чатов",
        chart_label_settings=chart_label_settings,
    )
    render_period_summary(
        project_id,
        project_name,
        selected_period_ids,
        messages,
        events_agg,
        periods,
        role,
        profile="driver_chats",
        metrics=metrics,
        branding=report_branding,
    )
    section_options = ["Клиентский обзор", "Инфоповоды", "Ключевые сообщения"]
    if len(selected_period_ids) >= 2:
        section_options.append("Динамика")
    section = st.radio("Раздел аналитики", section_options, horizontal=True, key="taxi_dashboard_section")

    if section == "Клиентский обзор":
        render_client_insights(messages, events_agg, periods, selected_period_ids, profile="driver_chats")
    elif section == "Инфоповоды":
        render_small_events_notice(hidden_events, hidden_messages, min_event_messages)
        render_events(project_id, role, events_agg, messages, manual_state)
    elif section == "Ключевые сообщения":
        render_messages_block(messages, project_id=project_id)
    elif section == "Динамика":
        render_period_dynamics(messages, periods, selected_period_ids)


def main() -> None:
    args = parse_args()
    reset_perf_events()
    st.set_page_config(page_title=APP_TITLE, layout="wide")
    st.title(APP_TITLE)
    st.caption(APP_VERSION)

    if not supabase_configured():
        st.error("Supabase не настроен. Добавьте SUPABASE_URL и SUPABASE_SERVICE_ROLE_KEY в Secrets.")
        st.stop()

    is_admin = is_platform_admin()
    project_id, role, projects = render_project_access(is_admin)

    if is_admin:
        with st.sidebar.expander("Управление платформой", expanded=False):
            if st.button("Открыть управление проектами"):
                st.session_state["platform_page"] = "projects"
    if project_id:
        page_options = ["Дашборд", "Загрузка файла", "История периодов"]
    else:
        page_options = []
    if is_admin:
        page_options.append("Проекты")
    if not page_options:
        st.info("Выберите проект или войдите как владелец платформы.")
        if is_admin:
            render_project_manager(projects)
        return
    page = st.sidebar.radio("Раздел", page_options, index=page_options.index("Проекты") if st.session_state.get("platform_page") == "projects" and "Проекты" in page_options else 0)

    if page == "Проекты":
        render_project_manager(projects)
        return

    if not project_id:
        st.info("Введите код доступа к проекту или войдите как владелец платформы.")
        return

    project_row = projects[projects["project_id"].astype(str) == str(project_id)]
    current_project_row = project_row.iloc[0] if not project_row.empty else None
    project_name = str(current_project_row.get("project_name") if current_project_row is not None else project_id)
    project_profile = project_topic_profile(current_project_row)
    current_project_settings = project_settings_from_row(current_project_row) if current_project_row is not None else {}
    chart_label_settings = chart_label_settings_from_project_settings(current_project_settings)
    report_branding = report_branding_from_project_settings(current_project_settings, project_name=project_name)
    st.sidebar.markdown(f"**Текущий проект:**  \n{project_name}")
    st.sidebar.caption(f"Профиль: {ALGORITHM_PROFILE_OPTIONS.get(project_profile, project_profile)}")
    if role == "admin":
        st.sidebar.checkbox("Диагностика скорости", value=False, key="platform_perf_debug")

    if page == "Загрузка файла":
        render_upload_page(project_id, role, args.work_dir)
        return
    if page == "История периодов":
        render_period_history(project_id, role)
        return

    selected_period_ids, periods = render_period_selector(project_id)
    if not selected_period_ids:
        st.info("Выберите период или загрузите первый файл.")
        return

    with st.spinner("Загружаю данные проекта..."):
        with perf_block("dashboard.load_generated_tables", project_id=project_id, periods=len(selected_period_ids)):
            events, discussions, messages, discussion_messages, event_discussions = load_generated_tables(project_id, selected_period_ids)
    with perf_block("dashboard.enrich_messages", project_id=project_id):
        enriched_messages = enrich_messages(messages, event_discussions, discussion_messages, events)
    with perf_block("dashboard.apply_manual_overrides", project_id=project_id):
        events, enriched_messages, manual_state = apply_manual_overrides(project_id, events, enriched_messages)
    with perf_block("dashboard.prepare_messages", project_id=project_id):
        enriched_messages = prepare_dashboard_messages(enriched_messages)
    render_perf_sidebar()

    if is_taxi_project_profile(project_profile):
        render_taxi_dashboard(
            project_id,
            project_name,
            role,
            selected_period_ids,
            periods,
            events,
            enriched_messages,
            manual_state,
            chart_label_settings=chart_label_settings,
            report_branding=report_branding,
        )
        return

    raw_events_agg = aggregate_events(events)
    min_event_messages = render_min_event_messages_control(project_profile, events, key="main_min_event_messages")
    events_agg, hidden_events, hidden_messages = filter_small_events(raw_events_agg, min_event_messages)

    # Brand Analytics projects must show only system tags from columns after
    # `Обработано`. This prevents legacy taxi/generic labels from appearing
    # in the tag block after algorithm updates.
    enriched_messages = clean_brand_analytics_tags(enriched_messages)
    enriched_messages = prepare_dashboard_messages(enriched_messages)

    metrics = render_project_intro(
        project_name,
        enriched_messages,
        periods,
        selected_period_ids,
        profile_label=ALGORITHM_PROFILE_OPTIONS.get(project_profile, project_profile),
        chart_label_settings=chart_label_settings,
    )
    render_period_summary(
        project_id,
        project_name,
        selected_period_ids,
        enriched_messages,
        events_agg,
        periods,
        role,
        profile=project_profile,
        metrics=metrics,
        branding=report_branding,
    )
    section_options = ["Клиентский обзор", "Теги", "Инфоповоды", "Ключевые сообщения"]
    if len(selected_period_ids) >= 2:
        section_options.append("Динамика")
    section = st.radio("Раздел аналитики", section_options, horizontal=True, key="main_dashboard_section")

    if section == "Клиентский обзор":
        render_client_insights(enriched_messages, events_agg, periods, selected_period_ids, profile=project_profile)
    elif section == "Теги":
        render_tag_statistics(enriched_messages, project_id=project_id)
    elif section == "Инфоповоды":
        render_small_events_notice(hidden_events, hidden_messages, min_event_messages)
        render_events(project_id, role, events_agg, enriched_messages, manual_state)
    elif section == "Ключевые сообщения":
        render_messages_block(enriched_messages, project_id=project_id)
    elif section == "Динамика":
        render_period_dynamics(enriched_messages, periods, selected_period_ids)


if __name__ == "__main__":
    main()
